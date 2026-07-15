from io import BytesIO
import zipfile

from django.contrib.auth.models import User
from django.test import TestCase
from docx import Document
from lxml import etree

from core.models import Answer, Definition, Question, QuestionRelationship, Reference


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W_NS}


class PaperExportTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="paper-user", password="pass")
        self.other_user = User.objects.create_user(username="other-user", password="pass")

        self.root = Question.objects.create(
            question_text="Büyük Başlık",
            user=self.user,
        )
        self.child = Question.objects.create(
            question_text="Alt Başlık",
            user=self.user,
        )
        self.grandchild = Question.objects.create(
            question_text="İkinci Alt Başlık",
            user=self.user,
        )
        QuestionRelationship.objects.create(
            parent=self.root,
            child=self.child,
            user=self.user,
        )
        QuestionRelationship.objects.create(
            parent=self.child,
            child=self.grandchild,
            user=self.user,
        )

        self.aksoy = Reference.objects.create(
            author_surname="Aksoy",
            author_name="Ayşe",
            year=2018,
            metin_ismi="Alfabetik Kaynak",
            rest="Ankara: Örnek Yayınları.",
            created_by=self.user,
        )
        self.zengin = Reference.objects.create(
            author_surname="Zengin; Kaya",
            author_name="Zeynep; Kemal",
            year=2020,
            metin_ismi="Çok Yazarlı Kaynak",
            rest="İstanbul: Akademi.",
            created_by=self.user,
        )
        self.oz = Reference.objects.create(
            author_surname="Öz",
            author_name="Özlem",
            year=2019,
            metin_ismi="Türkçe Alfabetik Sıra",
            rest="İzmir: Deneme.",
            created_by=self.user,
        )
        self.definition_source = Reference.objects.create(
            author_surname="Yalçın",
            author_name="Yasemin",
            year=2021,
            metin_ismi="Tanım Kaynağı",
            rest="Bursa: Kavram Yayınları.",
            created_by=self.user,
        )
        self.definition = Definition.objects.create(
            user=self.user,
            question=self.root,
            definition_text=(
                f"Tanım dipnotu (kaynak:{self.definition_source.id}, sayfa:31)."
            ),
        )

        self.root_answer = Answer.objects.create(
            question=self.root,
            user=self.user,
            answer_text=(
                f"**Temel sav** (k:{self.zengin.id} s:12). "
                f"-g- İlk dipnot (kaynak:{self.aksoy.id}, sayfa:44) -g-\n\n"
                "Biçimler: [Mavi bağlantı](https://example.com), **kalın**, "
                "*italik* ve ***kalın italik***. "
                f"(t:Temel kavram:{self.definition.id})\n\n"
                "-- Birinci Bölüm\n\n"
                "Birinci bölüm metni.\n"
                "---- İkinci Bölüm\n"
                "İkinci bölüm metni.\n"
                "------ Üçüncü Bölüm\n"
                "Üçüncü bölüm metni.\n\n"
                "1. Birinci düzey\n"
                "1.2. İkinci düzey\n"
                "1.2.1. Üçüncü düzey\n"
                "1.2.1.1. Dördüncü düzey"
            ),
        )
        self.child_answer = Answer.objects.create(
            question=self.child,
            user=self.user,
            answer_text="Alt bölüm --gizli--İkinci dipnot--gizli-- metni.",
        )
        self.grandchild_answer = Answer.objects.create(
            question=self.grandchild,
            user=self.user,
            answer_text=(
                "[Bağlantı](https://example.com) içeren son bölüm "
                f"(k:{self.oz.id})."
            ),
        )
        self.client.force_login(self.user)

    def download(self):
        return self.client.post(
            f"/profile/{self.user.username}/download_entries_paper/",
            {
                "entry_ids": ",".join(
                    str(answer.id)
                    for answer in (
                        self.root_answer,
                        self.child_answer,
                        self.grandchild_answer,
                    )
                ),
                "order": "oldest",
                "root_question_id": str(self.root.id),
            },
        )

    def test_paper_export_formats_headings_citations_and_bibliography(self):
        response = self.download()

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("paper-user_paper.docx", response["Content-Disposition"])

        document = Document(BytesIO(response.content))
        paragraphs = document.paragraphs
        combined_text = "\n".join(paragraph.text for paragraph in paragraphs)
        styles_by_text = {
            paragraph.text: paragraph.style.name
            for paragraph in paragraphs
            if paragraph.text in {
                self.root.question_text,
                self.child.question_text,
                self.grandchild.question_text,
            }
        }

        self.assertEqual(styles_by_text[self.root.question_text], "Heading 1")
        self.assertEqual(styles_by_text[self.child.question_text], "Heading 2")
        self.assertEqual(styles_by_text[self.grandchild.question_text], "Heading 3")
        content_heading_styles = {
            paragraph.text: paragraph.style.name
            for paragraph in paragraphs
            if paragraph.text in {
                "Birinci Bölüm",
                "İkinci Bölüm",
                "Üçüncü Bölüm",
            }
        }
        self.assertEqual(content_heading_styles["Birinci Bölüm"], "Heading 2")
        self.assertEqual(content_heading_styles["İkinci Bölüm"], "Heading 4")
        self.assertEqual(content_heading_styles["Üçüncü Bölüm"], "Heading 6")
        self.assertNotIn("-- Birinci Bölüm", combined_text)
        self.assertNotIn("---- İkinci Bölüm", combined_text)
        self.assertNotIn("------ Üçüncü Bölüm", combined_text)
        self.assertIn("(Zengin vd., 2020, s. 12)", combined_text)
        self.assertNotIn("(k:", combined_text)
        self.assertNotIn("(kaynak:", combined_text)
        self.assertNotIn("-g-", combined_text)
        self.assertNotIn("--gizli--", combined_text)
        self.assertNotIn("(t:Temel kavram:", combined_text)
        self.assertIn("Temel kavram", combined_text)
        self.assertNotIn(self.root_answer.created_at.strftime("%Y-%m-%d"), combined_text)

        bibliography_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text == "Kaynakça" and paragraph.style.name == "Heading 1"
        )
        bibliography_texts = [
            paragraph.text
            for paragraph in paragraphs[bibliography_index + 1:]
            if paragraph.text.strip()
        ]
        self.assertTrue(bibliography_texts[0].startswith("Aksoy, Ayşe (2018)."))
        self.assertTrue(bibliography_texts[1].startswith("Öz, Özlem (2019)."))
        self.assertTrue(bibliography_texts[2].startswith("Yalçın, Yasemin (2021)."))
        self.assertTrue(bibliography_texts[3].startswith("Zengin, Zeynep; Kaya, Kemal (2020)."))

    def test_paper_export_preserves_inline_formatting_and_blue_links(self):
        response = self.download()

        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            document_root = etree.fromstring(archive.read("word/document.xml"))

        link = document_root.xpath(
            ".//w:hyperlink[.//w:t='Mavi bağlantı']",
            namespaces=NS,
        )[0]
        self.assertIsNotNone(link.get(f"{{{R_NS}}}id"))
        self.assertEqual(
            link.xpath("string(.//w:color/@w:val)", namespaces=NS),
            "0563C1",
        )
        self.assertEqual(
            link.xpath("string(.//w:u/@w:val)", namespaces=NS),
            "single",
        )

        def runs_for_text(value):
            return document_root.xpath(
                f".//w:r[w:t={value!r}]",
                namespaces=NS,
            )

        bold_run = runs_for_text("kalın")[0]
        italic_run = runs_for_text("italik")[0]
        bold_italic_run = runs_for_text("kalın italik")[0]
        self.assertTrue(bold_run.xpath("./w:rPr/w:b", namespaces=NS))
        self.assertFalse(bold_run.xpath("./w:rPr/w:i", namespaces=NS))
        self.assertTrue(italic_run.xpath("./w:rPr/w:i", namespaces=NS))
        self.assertFalse(italic_run.xpath("./w:rPr/w:b", namespaces=NS))
        self.assertTrue(bold_italic_run.xpath("./w:rPr/w:b", namespaces=NS))
        self.assertTrue(bold_italic_run.xpath("./w:rPr/w:i", namespaces=NS))

    def test_paper_export_indents_body_paragraphs_and_four_outline_levels(self):
        response = self.download()
        document = Document(BytesIO(response.content))

        body_paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "Birinci bölüm metni."
        )
        self.assertEqual(body_paragraph.style.name, "Paper Body")
        self.assertAlmostEqual(
            body_paragraph.style.paragraph_format.first_line_indent.cm,
            1.25,
            places=2,
        )

        numbered_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "Paper Numbered Item"
        ]
        self.assertEqual(
            [paragraph.text for paragraph in numbered_paragraphs],
            [
                "1.\tBirinci düzey",
                "1.2.\tİkinci düzey",
                "1.2.1.\tÜçüncü düzey",
                "1.2.1.1.\tDördüncü düzey",
            ],
        )
        for paragraph, expected_indent in zip(
            numbered_paragraphs,
            (1.1, 1.85, 2.6, 3.35),
        ):
            self.assertAlmostEqual(
                paragraph.paragraph_format.left_indent.cm,
                expected_indent,
                places=2,
            )
            self.assertAlmostEqual(
                paragraph.paragraph_format.first_line_indent.cm,
                -1.1,
                places=2,
            )

    def test_paper_export_contains_toc_and_custom_star_footnotes(self):
        response = self.download()

        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            self.assertIn("word/footnotes.xml", archive.namelist())
            document_root = etree.fromstring(archive.read("word/document.xml"))
            footnotes_root = etree.fromstring(archive.read("word/footnotes.xml"))

        references = document_root.xpath(".//w:footnoteReference", namespaces=NS)
        self.assertEqual(len(references), 3)
        self.assertTrue(
            all(
                reference.get(f"{{{W_NS}}}customMarkFollows") == "1"
                for reference in references
            )
        )
        body_text_nodes = document_root.xpath(".//w:body//w:t/text()", namespaces=NS)
        self.assertEqual(
            [
                "".join(reference.getparent().xpath("./w:t/text()", namespaces=NS))
                for reference in references
            ],
            ["*", "**", "***"],
        )
        self.assertNotIn("İçindekiler hazırlanıyor.", body_text_nodes)
        self.assertEqual(
            len(document_root.xpath(".//w:bookmarkStart", namespaces=NS)),
            7,
        )
        self.assertEqual(
            len(document_root.xpath(".//w:hyperlink[@w:anchor]", namespaces=NS)),
            7,
        )

        toc_text = " ".join(
            document_root.xpath(
                ".//w:body/w:p[w:pPr/w:pStyle[@w:val='PaperTOCEntry']]//w:t/text()",
                namespaces=NS,
            )
        )
        self.assertIn("Birinci Bölüm", toc_text)
        self.assertIn("İkinci Bölüm", toc_text)
        self.assertIn("Üçüncü Bölüm", toc_text)

        note_text = " ".join(
            footnote.text
            for footnote in footnotes_root.xpath(
                ".//w:footnote[number(@w:id) > 0]//w:t",
                namespaces=NS,
            )
            if footnote.text
        )
        self.assertIn("İlk dipnot (Aksoy, 2018, s. 44)", note_text)
        self.assertIn("Tanım dipnotu (Yalçın, 2021, s. 31)", note_text)
        self.assertIn("İkinci dipnot", note_text)
        self.assertEqual(
            footnotes_root.xpath(
                ".//w:footnote[number(@w:id) > 0]"
                "//w:r[w:rPr/w:rStyle[@w:val='FootnoteReference']]"
                "/w:t/text()",
                namespaces=NS,
            ),
            ["*", "**", "***"],
        )
        self.assertFalse(
            footnotes_root.xpath(
                ".//w:footnote[number(@w:id) > 0]//w:footnoteRef",
                namespaces=NS,
            )
        )

    def test_paper_export_rejects_another_user(self):
        self.client.force_login(self.other_user)

        response = self.client.post(
            f"/profile/{self.user.username}/download_entries_paper/",
        )

        self.assertEqual(response.status_code, 403)

    def test_download_modal_offers_paper_format(self):
        response = self.client.get(f"/profile/{self.user.username}/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'value="paper"')
        self.assertContains(response, "download_entries_paper")
