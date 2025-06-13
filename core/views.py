import json
import random
import re
import colorsys
from collections import Counter, defaultdict
from datetime import timedelta
from io import BytesIO
from urllib.parse import unquote
from .models import IATResult, Kenarda
# from .models import Kenarda
from django.conf import settings
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.contenttypes.models import ContentType
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.core.serializers import serialize
from django.db import transaction
from django.db.models import Q, Count, Max, F, Sum, Prefetch
from django.db.models.functions import Coalesce, Lower
from django.http import (
    HttpResponse,
    JsonResponse,
    HttpResponseForbidden
)

from django.shortcuts import (
    render,
    redirect,
    get_object_or_404,
)
from django.template.loader import render_to_string
from django.urls import reverse
from django.utils import timezone
from django.utils.timezone import now
from django.views.decorators.cache import cache_control
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST

# 3. parti
from openpyxl import Workbook
from openpyxl.styles import Font
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Uygulama içi
from .models import (
    Poll, Reference, PollOption, RandomSentence, PollVote,
    Invitation, UserProfile, Question, Answer, StartingQuestion,
    Vote, Message, SavedItem, PinnedEntry, Definition, CikisTesti, CikisTestiSoru, CikisTestiSik,CikisTestiResult
)
from .forms import (
    InvitationForm, AnswerForm, StartingQuestionForm, SignupForm, LoginForm, 
    QuestionForm, ProfilePhotoForm, MessageForm, DefinitionForm, 
    RandomSentenceForm, PollForm, ReferenceForm,CikisTestiForm, CikisTestiSoruForm, CikisTestiSikForm
)




@login_required
def get_user_questions(request):
    q = request.GET.get('q', '').strip()
    # İstediğiniz filtreleme: örneğin sadece kendi sorularınız veya profildeki tüm sorular
    questions = Question.objects.filter(user=request.user)
    if q:
        questions = questions.filter(question_text__icontains=q)
    data = []
    for question in questions:
        data.append({
            'id': question.id,
            'question_text': question.question_text,
            'detail_url': reverse('question_detail', args=[question.id]),
        })
    return JsonResponse({'questions': data})

@login_required
def get_user_answers(request):
    q = request.GET.get('q', '').strip()
    answers = Answer.objects.filter(user=request.user)
    if q:
        answers = answers.filter(
            Q(answer_text__icontains=q) | Q(question__question_text__icontains=q)
        )
    data = []
    for answer in answers:
        data.append({
            'id': answer.id,
            'question_text': answer.question.question_text,  # Bunu da gönder
            'answer_text': answer.answer_text[:80] + '...' if len(answer.answer_text) > 80 else answer.answer_text,
            'detail_url': reverse('single_answer', args=[answer.question.id, answer.id]),
            'question_url': reverse('question_detail', args=[answer.question.id]),
            'created_at': answer.created_at.strftime("%d %b %Y %H:%M"),
        })
    return JsonResponse({'answers': data})

@login_required
def get_saved_items(request):
    q = request.GET.get('q', '').strip()
    saved_items = SavedItem.objects.filter(user=request.user)
    filtered_items = []
    for item in saved_items:
        instance = item.content_type.get_object_for_this_type(id=item.object_id)
        if item.content_type.model == 'question':
            text = instance.question_text
            # Soru için sadece başlık aranır
            matches = not q or q.lower() in text.lower()
        elif item.content_type.model == 'answer':
            # Hem yanıt hem başlık aranır
            text = instance.answer_text
            matches = (
                not q
                or q.lower() in text.lower()
                or q.lower() in instance.question.question_text.lower()
            )
        else:
            text = ''
            matches = False
        if matches:
            filtered_items.append({
                'type': item.content_type.model,
                'id': instance.id,
                'text': text[:80] + '...' if len(text) > 80 else text,
                'detail_url': reverse('question_detail', args=[instance.id]) if item.content_type.model == 'question' else reverse('single_answer', args=[instance.question.id, instance.id]),
            })
    return JsonResponse({'saved_items': filtered_items})



def signup(request):
    if request.method == 'POST':
        form = SignupForm(request.POST)
        if form.is_valid():
            invitation_code = form.cleaned_data['invitation_code']
            # Davetiye kodunu kontrol et
            try:
                invitation = Invitation.objects.get(code=invitation_code, is_used=False)
            except Invitation.DoesNotExist:
                messages.error(request, 'Geçersiz veya kullanılmış davet kodu.')
                return render(request, 'core/signup.html', {'form': form})

            user = form.save()  # Kullanıcı oluşturuldu

            # UserProfile'ı kontrol edip oluşturun, eğer yoksa
            try:
                user_profile = user.userprofile
            except UserProfile.DoesNotExist:
                if user.is_superuser:
                    quota = 999999999
                else:
                    quota = 0
                user_profile = UserProfile.objects.create(user=user, invitation_quota=quota)

            # Davetiye kodunu kullanılmış olarak işaretle ve quota güncelle
            invitation.is_used = True
            invitation.used_by = user
            invitation.save()

            user_profile.invitation_quota += invitation.quota_granted
            user_profile.save()

            login(request, user)
            return redirect('user_homepage')
        else:
            # Form geçersizse hataları göster
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")
    else:
        form = SignupForm()

    return render(request, 'core/signup.html', {'form': form})


def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('user_homepage')
    else:
        form = LoginForm()
    return render(request, 'core/login.html', {'form': form})

def user_logout(request):
    logout(request)
    return redirect('user_homepage')

@login_required
def send_invitation(request):
    user_profile = request.user.userprofile
    if request.method == 'POST':
        form = InvitationForm(request.POST)
        if form.is_valid():
            quota_granted = form.cleaned_data['quota_granted']
            if user_profile.invitation_quota >= quota_granted:
                with transaction.atomic():
                    # Davet kodunu oluştur
                    invitation = form.save(commit=False)
                    invitation.sender = request.user
                    invitation.quota_granted = quota_granted
                    invitation.save()

                    # Kullanıcının davet hakkını düşür
                    user_profile.invitation_quota -= quota_granted
                    user_profile.save()

                # Oluşturulan kodu ve güncel davet hakkını şablona aktar
                return render(request, 'core/send_invitation.html', {
                    'form': InvitationForm(),  # Yeni davetler için boş form
                    'invitation_code': invitation.code,
                    'quota_granted': quota_granted,
                    'invitation_quota': user_profile.invitation_quota,
                })
            else:
                messages.error(request, 'Yeterli davet hakkınız yok.')
    else:
        form = InvitationForm()

    return render(request, 'core/send_invitation.html', {
        'form': form,
        'invitation_quota': user_profile.invitation_quota,
    })

@login_required
def profile(request):
    user = request.user
    user_profile = user.userprofile

    # Kullanıcının soruları ve yanıtları
    questions = Question.objects.filter(user=user)
    answers = Answer.objects.filter(user=user)

    # Takipçi ve takip edilen sayıları
    follower_count = user_profile.followers.count()
    following_count = user_profile.following.count()

    # Sabitlenmiş giriş (pinned_entry)
    try:
        pinned_entry = PinnedEntry.objects.get(user=user)
    except PinnedEntry.DoesNotExist:
        pinned_entry = None

    # En çok kullanılan kelimeler
    top_words = get_top_words(user)

    context = {
        'profile_user': user,
        'user_profile': user_profile,
        'questions': questions,
        'answers': answers,
        'follower_count': follower_count,
        'following_count': following_count,
        'pinned_entry': pinned_entry,
        'top_words': top_words,
    }
    return redirect('user_profile', username=request.user.username)


@login_required
def user_profile(request, username):
    profile_user = get_object_or_404(User, username=username)
    user_profile = profile_user.userprofile
    active_tab = request.GET.get('tab', 'girdiler')
    is_own_profile = (request.user == profile_user)

    # Her sekme için context anahtarları baştan boş atanıyor
    context = {
        'profile_user': profile_user,
        'user_profile': user_profile,
        'is_own_profile': is_own_profile,
        'active_tab': active_tab,
        'answers': None,
        'definitions_page': None,
        'references_page': None,
        'saved_items_page': None,
        'top_words': None,
        'exclude_words': None,
        'exclude_words_list': None,
        'search_word': None,
        'search_word_count': None,
        # ... diğer anahtarlar da buraya eklenecekse boş bırak
    }

    # Takip ve takipçi sayısı
    context['follower_count'] = user_profile.followers.count()
    context['following_count'] = user_profile.following.count()
    context['is_following'] = (
        request.user.is_authenticated and
        request.user != profile_user and
        request.user.userprofile.following.filter(user=profile_user).exists()
    )

    # Sabitlenmiş entry
    try:
        context['pinned_entry'] = PinnedEntry.objects.select_related('answer__question').get(user=profile_user)
    except PinnedEntry.DoesNotExist:
        context['pinned_entry'] = None

    # Hangi sekmedeyse sadece oradaki veriyi doldur
    if active_tab == 'girdiler':
        answers_list = Answer.objects.filter(user=profile_user).select_related('question', 'user').order_by('-created_at')
        answer_paginator = Paginator(answers_list, 10)
        answer_page = request.GET.get('answer_page', 1)
        try:
            context['answers'] = answer_paginator.page(answer_page)
        except (PageNotAnInteger, EmptyPage):
            context['answers'] = answer_paginator.page(1)

    elif active_tab == 'tanimlar':
        definitions_qs = Definition.objects.filter(user=profile_user).select_related('question')
        search_query = request.GET.get('q', '').strip()
        if search_query:
            definitions_qs = definitions_qs.filter(
                Q(definition_text__icontains=search_query) |
                Q(question__question_text__icontains=search_query)
            )
        def_paginator = Paginator(definitions_qs, 5)
        def_page = request.GET.get('d_page', 1)
        try:
            context['definitions_page'] = def_paginator.page(def_page)
        except (PageNotAnInteger, EmptyPage):
            context['definitions_page'] = def_paginator.page(1)

    elif active_tab == 'kaynaklarim':
        user_references = Reference.objects.filter(created_by=profile_user)
        search_query = request.GET.get('q', '').strip()
        if search_query:
            user_references = user_references.filter(
                Q(author_surname__icontains=search_query) |
                Q(author_name__icontains=search_query) |
                Q(rest__icontains=search_query) |
                Q(abbreviation__icontains=search_query)
            )
        ref_paginator = Paginator(user_references, 5)
        r_page = request.GET.get('r_page', 1)
        try:
            context['references_page'] = ref_paginator.page(r_page)
        except (PageNotAnInteger, EmptyPage):
            context['references_page'] = ref_paginator.page(1)

    elif active_tab == 'kaydedilenler' and is_own_profile:
        user_saved = SavedItem.objects.filter(user=profile_user).select_related('content_type').order_by('-saved_at')
        saved_items_list = []
        content_type_map = {ct.model: ct for ct in ContentType.objects.all()}
        question_ids = [si.object_id for si in user_saved if si.content_type.model == 'question']
        answer_ids = [si.object_id for si in user_saved if si.content_type.model == 'answer']
        question_map = {q.id: q for q in Question.objects.filter(id__in=question_ids)}
        answer_map = {a.id: a for a in Answer.objects.select_related('question').filter(id__in=answer_ids)}
        for item in user_saved:
            if item.content_type.model == 'question':
                obj = question_map.get(item.object_id)
                if obj:
                    saved_items_list.append({'type': 'question', 'object': obj})
            elif item.content_type.model == 'answer':
                obj = answer_map.get(item.object_id)
                if obj:
                    saved_items_list.append({'type': 'answer', 'object': obj})
        s_page = request.GET.get('s_page', 1)
        saved_paginator = Paginator(saved_items_list, 5)
        try:
            context['saved_items_page'] = saved_paginator.page(s_page)
        except (PageNotAnInteger, EmptyPage):
            context['saved_items_page'] = saved_paginator.page(1)

    elif active_tab in ['istatistikler', 'kelimeler']:
        questions_list = Question.objects.filter(user=profile_user)
        answers_list = Answer.objects.filter(user=profile_user)
        question_texts = questions_list.values_list('question_text', flat=True)
        answer_texts = answers_list.values_list('answer_text', flat=True)
        all_texts = (' '.join(question_texts) + ' ' + ' '.join(answer_texts)).lower()
        words = re.findall(r'\b\w+\b', all_texts)
        exclude_words_str = request.GET.get('exclude_words', '')
        exclude_words_list = [w.strip() for w in exclude_words_str.split(',') if w.strip()]
        exclude_words_set = set(word.lower() for word in exclude_words_list)
        exclude_word = request.GET.get('exclude_word', '').strip().lower()
        if exclude_word:
            exclude_words_set.add(exclude_word)
        include_word = request.GET.get('include_word', '').strip().lower()
        if include_word and include_word in exclude_words_set:
            exclude_words_set.remove(include_word)
        exclude_words_list = sorted(list(exclude_words_set))
        exclude_words_str = ', '.join(exclude_words_list)
        filtered_words = [word for word in words if word not in exclude_words_set]
        word_counts = Counter(filtered_words)
        top_words = word_counts.most_common(20)
        search_word = request.GET.get('search_word', '').strip().lower()
        search_word_count = None
        if search_word:
            search_word_count = word_counts.get(search_word, 0)
        context.update({
            'top_words': top_words,
            'exclude_words': exclude_words_str,
            'exclude_words_list': exclude_words_list,
            'search_word': search_word,
            'search_word_count': search_word_count,
        })
        # İstatistikler için:
        total_upvotes = ((questions_list.aggregate(total=Sum('upvotes'))['total'] or 0) +
                         (answers_list.aggregate(total=Sum('upvotes'))['total'] or 0))
        total_downvotes = ((questions_list.aggregate(total=Sum('downvotes'))['total'] or 0) +
                           (answers_list.aggregate(total=Sum('downvotes'))['total'] or 0))
        content_type_question = ContentType.objects.get_for_model(Question)
        content_type_answer = ContentType.objects.get_for_model(Answer)
        total_saves_questions = SavedItem.objects.filter(
            content_type=content_type_question, object_id__in=questions_list
        ).count()
        total_saves_answers = SavedItem.objects.filter(
            content_type=content_type_answer, object_id__in=answers_list
        ).count()
        total_saves = total_saves_questions + total_saves_answers
        most_upvoted_question = questions_list.order_by('-upvotes').first()
        most_upvoted_answer = answers_list.order_by('-upvotes').first()
        most_upvoted_entry = max(
            (e for e in [most_upvoted_question, most_upvoted_answer] if e),
            key=lambda x: x.upvotes,
            default=None
        )
        most_downvoted_question = questions_list.order_by('-downvotes').first()
        most_downvoted_answer = answers_list.order_by('-downvotes').first()
        most_downvoted_entry = max(
            (e for e in [most_downvoted_question, most_downvoted_answer] if e),
            key=lambda x: x.downvotes,
            default=None
        )
        most_saved_question = questions_list.annotate(save_count=Count('saveditem')).order_by('-save_count').first()
        most_saved_answer = answers_list.annotate(save_count=Count('saveditem')).order_by('-save_count').first()
        most_saved_entry = max(
            (e for e in [most_saved_question, most_saved_answer] if e),
            key=lambda x: getattr(x, 'save_count', 0),
            default=None
        )
        context.update({
            'total_upvotes': total_upvotes,
            'total_downvotes': total_downvotes,
            'total_saves': total_saves,
            'most_upvoted_entry': most_upvoted_entry,
            'most_downvoted_entry': most_downvoted_entry,
            'most_saved_entry': most_saved_entry,
        })

    if active_tab == 'davetler' and is_own_profile:
        invitations = Invitation.objects.filter(sender=request.user).order_by('-created_at')
        total_invitations = invitations.count()
        used_invitations = invitations.filter(is_used=True).count()
        remaining_invitations = user_profile.invitation_quota - total_invitations
        context.update({
            'invitations': invitations,
            'total_invitations': total_invitations,
            'used_invitations': used_invitations,
            'remaining_invitations': remaining_invitations,
        })

    if active_tab == 'davet_aagac' and is_own_profile:
        context['invitation_tree'] = get_invitation_tree(request.user)

    return render(request, 'core/user_profile.html', context)

@login_required
def create_invitation(request):
    if request.method == 'POST':
        user_profile = request.user.userprofile
        total_invitations = Invitation.objects.filter(sender=request.user).count()
        remaining_invitations = user_profile.invitation_quota - total_invitations

        quota_granted = int(request.POST.get('quota_granted', 1))

        if quota_granted > remaining_invitations:
            messages.error(request, f"En fazla {remaining_invitations} adet davetiye kotası verebilirsiniz.")
        else:
            Invitation.objects.create(sender=request.user, quota_granted=quota_granted)
            # Kullanıcının davetiye kotasını azalt
            user_profile.invitation_quota -= quota_granted
            user_profile.save()
            messages.success(request, f"Davetiye oluşturuldu. Kullanıcıya {quota_granted} adet davetiye kotası verilecek.")

    # Yönlendirme yaparken 'tab' parametresini ekleyin
    return redirect(f'{reverse("user_profile", args=[request.user.username])}?tab=davetler')


def get_invitation_tree(user):
    invitations = Invitation.objects.filter(sender=user)
    tree = []
    for invite in invitations:
        if invite.is_used and invite.used_by:
            invited_user = invite.used_by
            subtree = get_invitation_tree(invited_user)
            tree.append({'user': invited_user, 'children': subtree})
        else:
            tree.append({'code': invite.code, 'children': []})
    return tree


@login_required
def question_detail(request, question_id):
    # Soldaki tüm sorular ve pagination
    all_questions = get_today_questions_queryset()
    q_page_number = request.GET.get('q_page', 1)
    q_paginator = Paginator(all_questions, 20)
    all_questions_page = q_paginator.get_page(q_page_number)

    # Soru nesnesini al
    question = get_object_or_404(Question, id=question_id)

    # Filtre parametreleri
    my_answers = request.GET.get('my_answers')
    followed = request.GET.get('followed')
    username = request.GET.get('username', '').strip()
    keyword  = request.GET.get('keyword', '').strip()

    all_answers = question.answers.all().order_by('created_at')
    if my_answers == 'on':
        all_answers = all_answers.filter(user=request.user)
    if followed == 'on':
        followed_users = request.user.userprofile.following.all()
        all_answers = all_answers.filter(user__in=followed_users)
    if username:
        all_answers = all_answers.filter(user__username__iexact=username)
    if keyword:
        all_answers = all_answers.filter(answer_text__icontains=keyword)

    a_page_number = request.GET.get('a_page', 1)
    a_paginator = Paginator(all_answers, 10)
    answers_page = a_paginator.get_page(a_page_number)

    if request.method == 'POST':
        form = AnswerForm(request.POST)
        if form.is_valid():
            new_answer = form.save(commit=False)
            new_answer.user = request.user
            new_answer.question = question
            new_answer.save()
            Kenarda.objects.filter(user=request.user, question=question, is_sent=False).delete()
            return redirect('question_detail', question_id=question.id)
    else:
        form = AnswerForm()

    content_type_question = ContentType.objects.get_for_model(Question)
    user_has_saved_question = SavedItem.objects.filter(
        user=request.user,
        content_type=content_type_question,
        object_id=question.id
    ).exists()
    question_save_count = SavedItem.objects.filter(
        content_type=content_type_question,
        object_id=question.id
    ).count()

    content_type_answer = ContentType.objects.get_for_model(Answer)
    page_answer_ids = [ans.id for ans in answers_page]
    user_votes = Vote.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=page_answer_ids
    ).values('object_id', 'value')
    user_vote_dict = {vote['object_id']: vote['value'] for vote in user_votes}
    for ans in answers_page:
        ans.user_vote_value = user_vote_dict.get(ans.id, 0)
    saved_answer_ids = SavedItem.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=page_answer_ids
    ).values_list('object_id', flat=True)
    answer_save_counts = SavedItem.objects.filter(
        content_type=content_type_answer,
        object_id__in=page_answer_ids
    ).values('object_id').annotate(count=Count('id'))
    answer_save_dict = {item['object_id']: item['count'] for item in answer_save_counts}

    starting_question_ids = set(StartingQuestion.objects.values_list('question_id', flat=True))
    is_on_map = (question.id in starting_question_ids) or question.parent_questions.exists()

    context = {
        'question': question,
        'form': form,
        'all_questions_page': all_questions_page,
        'answers_page': answers_page,
        'user_has_saved_question': user_has_saved_question,
        'question_save_count': question_save_count,
        'saved_answer_ids': list(saved_answer_ids),
        'answer_save_dict': answer_save_dict,
        'my_answers': my_answers,
        'followed': followed,
        'filter_username': username,
        'filter_keyword': keyword,
        'is_on_map': is_on_map,
    }
    return render(request, 'core/question_detail.html', context)


@login_required
def add_answer(request, question_id):
    question = get_object_or_404(Question, id=question_id)
    if request.method == 'POST':
        form = AnswerForm(request.POST)
        if form.is_valid():
            answer = form.save(commit=False)
            answer.question = question
            answer.user = request.user
            answer.save()
            from .models import Kenarda
            silinen = Kenarda.objects.filter(user=request.user, question=question, is_sent=False)
            print("SİLİNECEK:", list(silinen))
            deleted_count, _ = silinen.delete()
            print("Kaç tane silindi:", deleted_count)
            return redirect('question_detail', question_id=question.id)
    else:
        form = AnswerForm()
    return render(request, 'core/add_answer.html', {'form': form, 'question': question})


@login_required
def sent_messages(request):
    messages = Message.objects.filter(sender=request.user).order_by('-timestamp')
    return render(request, 'core/sent_messages.html', {'messages': messages})

@login_required
def view_message(request, message_id):
    message = get_object_or_404(Message, id=message_id)
    if message.recipient != request.user and message.sender != request.user:
        return redirect('inbox')
    if message.recipient == request.user:
        message.is_read = True
        message.save()
    return render(request, 'core/view_message.html', {'message': message})


@login_required
def user_list(request):
    users = User.objects.exclude(id=request.user.id) \
                        .annotate(username_lower=Lower('username')) \
                        .order_by('username_lower')
    paginator = Paginator(users, 36)  # Sayfa başına 36 kullanıcı
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'core/user_list.html', {'users': page_obj})

@login_required
def follow_user(request, username):
    target_user = get_object_or_404(User, username=username)
    user_profile = request.user.userprofile
    target_profile = target_user.userprofile
    user_profile.following.add(target_profile)
    return redirect('user_profile', username=username)

@login_required
def unfollow_user(request, username):
    target_user = get_object_or_404(User, username=username)
    user_profile = request.user.userprofile
    target_profile = target_user.userprofile
    user_profile.following.remove(target_profile)
    return redirect('user_profile', username=username)

def search_suggestions(request):
    query = request.GET.get('q', '')
    suggestions = []

    # Kullanıcıları ara
    users = User.objects.filter(username__icontains=query)
    for user in users:
        suggestions.append({
            'label': '@' + user.username,
            'url': reverse('user_profile', args=[user.username])  # Burada user.username kullanılıyor
        })

    # Soruları ara
    questions = Question.objects.filter(question_text__icontains=query)
    for question in questions:
        suggestions.append({
            'label': question.question_text,
            'url': reverse('question_detail', args=[question.id])
        })

    # Gelen veriyi sunucu konsoluna yazdırın

    print('search_suggestions verisi:', json.dumps({'suggestions': suggestions}, indent=4))

    return JsonResponse({'suggestions': suggestions})

@login_required
def search(request):
    """
    Hem AJAX (autocomplete) isteklerini hem de gelişmiş arama parametrelerini destekleyen arama view'i.
    """

    # 1) AJAX isteği kontrolü
    is_ajax = (
        request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest' or
        request.GET.get('ajax') == '1'
    )

    if is_ajax:
        query = request.GET.get('q', '').strip()
        if query:
            questions = Question.objects.filter(question_text__icontains=query)
            users = User.objects.filter(username__icontains=query)

            results = []
            for q_obj in questions:
                results.append({
                    'type': 'question',
                    'id': q_obj.id,
                    'text': q_obj.question_text,
                    'url': reverse('question_detail', args=[q_obj.id]),
                })
            for user in users:
                results.append({
                    'type': 'user',
                    'id': user.id,
                    'username': user.username,
                    'text': '@' + user.username,
                    'url': reverse('user_profile', args=[user.username]),
                })

            return JsonResponse({'results': results})
        else:
            return JsonResponse({'results': []})

    # 2) Gelişmiş Arama Parametreleri
    username = request.GET.get('username', '').strip()
    date_from = request.GET.get('date_from', '').strip()
    date_to = request.GET.get('date_to', '').strip()
    keywords = request.GET.get('keywords', '').strip()
    search_in = request.GET.get('search_in', 'all')
    q_param = request.GET.get('q', '').strip()

    # 3) Başlangıç querysetleri
    questions = Question.objects.all()
    answers = Answer.objects.all()
    users_found = User.objects.none()

    # 4) Basit 'q' araması
    if q_param:
        questions = questions.filter(question_text__icontains=q_param)
        answers = answers.filter(answer_text__icontains=q_param)
        users_found = User.objects.filter(username__icontains=q_param)

    # 5) Gelişmiş parametrelerle filtreleme
    if username:
        questions = questions.filter(user__username__icontains=username)
        answers = answers.filter(user__username__icontains=username)

    if date_from:
        questions = questions.filter(created_at__date__gte=date_from)
        answers = answers.filter(created_at__date__gte=date_from)

    if date_to:
        questions = questions.filter(created_at__date__lte=date_to)
        answers = answers.filter(created_at__date__lte=date_to)

    if keywords:
        questions = questions.filter(question_text__icontains=keywords)
        answers = answers.filter(answer_text__icontains=keywords)

    # 6) Hangi tabloda arama yapılacağına karar ver
    if search_in == 'question':
        answers = Answer.objects.none()
    elif search_in == 'answer':
        questions = Question.objects.none()

    # 7) Soruları ve yanıtları birleştir
    combined_results = []
    for q in questions:
        combined_results.append({
            "type": "question",
            "object": q,
            "created_at": q.created_at
        })
    for a in answers:
        combined_results.append({
            "type": "answer",
            "object": a,
            "created_at": a.created_at
        })

    combined_results.sort(key=lambda x: x['created_at'], reverse=True)

    # 8) Pagination for combined results
    page_number = request.GET.get('page', 1)
    paginator = Paginator(combined_results, 15)
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    # 9) Kullanıcılar için pagination
    users_page_number = request.GET.get('users_page', 1)
    users_paginator = Paginator(users_found.order_by('username'), 10)
    try:
        users_paginated = users_paginator.page(users_page_number)
    except PageNotAnInteger:
        users_paginated = users_paginator.page(1)
    except EmptyPage:
        users_paginated = users_paginator.page(users_paginator.num_pages)

    # 10) Sonuçları template'e gönder
    context = {
        'results': page_obj,         # Birleşik sonuçlar burada!
        'users': users_paginated,    # Kullanıcılar ayrı.
        'query': q_param,
        'username': username,
        'date_from': date_from,
        'date_to': date_to,
        'keywords': keywords,
        'search_in': search_in,
        'page_obj': page_obj,        # Pagination için
    }

    return render(request, 'core/search_results.html', context)


@login_required
def add_question_from_search(request):
    all_questions = get_today_questions_queryset()
    query = request.GET.get('q', '').strip()

    if request.method == 'POST':
        answer_form = AnswerForm(request.POST)
        if answer_form.is_valid():
            # Yeni soru oluştur
            question = Question.objects.create(
                question_text=query,
                user=request.user,
                from_search=True  # Eğer modelinizde bu alan varsa
            )
            # Kullanıcıyı soru ile ilişkilendir
            question.users.add(request.user)
            
            # Yeni yanıt oluştur
            answer = answer_form.save(commit=False)
            answer.user = request.user
            answer.question = question
            answer.save()
            return redirect('question_detail', question_id=question.id)
    else:
        answer_form = AnswerForm()
    return render(request, 'core/add_question_from_search.html', {
        'query': query,
        'answer_form': answer_form,
        'all_questions': all_questions,
    })

def get_user_id(request, username):
    try:
        user = User.objects.get(username=username)
        return JsonResponse({'user_id': user.id})
    except User.DoesNotExist:
        return JsonResponse({'error': 'User not found'}, status=404)
    
@login_required
def add_question(request):
    if request.method == 'POST':
        form = QuestionForm(request.POST)
        if form.is_valid():
            question_text = form.cleaned_data['question_text']
            # Aynı soru metni varsa mevcut olanı kullan
            question, created = Question.objects.get_or_create(
                question_text=question_text,
                defaults={'user': request.user}
            )
            question.users.add(request.user)
            question.save()
            # Başlangıç sorusu olarak ekle
            StartingQuestion.objects.create(user=request.user, question=question)
            # Yanıtı kaydet
            Answer.objects.create(
                question=question,
                user=request.user,
                answer_text=form.cleaned_data.get('answer_text', '')
            )
            return redirect('user_homepage')
    else:
        form = QuestionForm()
    return render(request, 'core/add_question.html', {'form': form})

@login_required
def add_subquestion(request, question_id):
    parent_question = get_object_or_404(Question, id=question_id)
    all_questions = get_today_questions_queryset()
    if request.method == 'POST':
        form = QuestionForm(request.POST)
        if form.is_valid():
            subquestion_text = form.cleaned_data['question_text']
            answer_text = form.cleaned_data.get('answer_text', '')
            # Yeni veya mevcut alt soruyu oluştururken 'user' bilgisini ekliyoruz
            subquestion, created = Question.objects.get_or_create(
                question_text=subquestion_text,
                defaults={'user': request.user}
            )
            subquestion.users.add(request.user)
            parent_question.subquestions.add(subquestion)
            # Yanıtı kaydet
            Answer.objects.create(
                question=subquestion,
                user=request.user,
                answer_text=answer_text
            )
            messages.success(request, 'Alt soru başarıyla eklendi.')
            return redirect('question_detail', question_id=subquestion.id)
    else:
        form = QuestionForm()
    context = {
        'form': form,
        'parent_question': parent_question,
        'all_questions': all_questions,
    }
    return render(request, 'core/add_subquestion.html', context)

@login_required
def question_map(request):
    question_id = request.GET.get('question_id', None)
    # Başlangıç soruları
    starting_question_ids = StartingQuestion.objects.values_list('question_id', flat=True)
    # Sadece başlangıç soruları veya alt soruları getir
    questions = Question.objects.filter(
        Q(id__in=starting_question_ids) | Q(parent_questions__isnull=False)
    )

    nodes = {}
    links = []
    question_text_to_ids = defaultdict(list)

    # Build nodes dictionary keyed by question_text
    for question in questions:
        key = question.question_text
        question_text_to_ids[key].append(question.id)
        if key not in nodes:
            associated_users = list(question.users.all())
            user_ids = [user.id for user in associated_users]
            node = {
                "id": f"q{hash(key)}",  # Unique ID based on question_text
                "label": question.question_text,
                "users": user_ids,
                "size": 20 + 10 * (len(user_ids) - 1),
                "color": '',
                "question_id": question.id,  # Store a valid question ID
                "question_ids": [question.id],  # List of question IDs with same text
            }
            # Assign color based on user IDs
            if len(user_ids) == 1:
                node["color"] = get_user_color(user_ids[0])
            elif len(user_ids) > 1:
                node["color"] = '#CCCCCC'  # Grey for multiple users
            else:
                node["color"] = '#000000'  # Black if no user
            nodes[key] = node
        else:
            # Merge user IDs and update size
            existing_node = nodes[key]
            new_user_ids = [user.id for user in question.users.all()]
            combined_user_ids = list(set(existing_node["users"] + new_user_ids))
            existing_node["users"] = combined_user_ids
            existing_node["size"] = 20 + 5 * (len(combined_user_ids) - 1)
            existing_node["question_ids"].append(question.id)
            # Update color
            if len(combined_user_ids) == 1:
                existing_node["color"] = get_user_color(combined_user_ids[0])
            elif len(combined_user_ids) > 1:
                existing_node["color"] = '#CCCCCC'
            else:
                existing_node["color"] = '#000000'

    # Build links using question_text as keys
    link_set = set()
    for question in questions:
        source_key = question.question_text
        for subquestion in question.subquestions.all():
            target_key = subquestion.question_text
            if target_key in nodes:
                link_id = (nodes[source_key]["id"], nodes[target_key]["id"])
                if link_id not in link_set:
                    links.append({
                        "source": nodes[source_key]["id"],
                        "target": nodes[target_key]["id"]
                    })
                    link_set.add(link_id)

    question_nodes = {
        "nodes": list(nodes.values()),
        "links": links
    }
    return render(request, 'core/question_map.html', {
        'question_nodes': json.dumps(question_nodes),
        'focus_question_id': question_id,
    })

def get_user_color(user_id):
    hue = (user_id * 137.508) % 360  # Altın açı
    rgb = colorsys.hsv_to_rgb(hue / 360, 0.5, 0.95)
    hex_color = '#%02x%02x%02x' % (int(rgb[0]*255), int(rgb[1]*255), int(rgb[2]*255))
    return hex_color

def about(request):
    return render(request, 'core/about.html')

@login_required
def user_settings(request):
    profile, created = UserProfile.objects.get_or_create(user=request.user)
    if request.method == 'POST':
        if 'reset' in request.POST:
            # Varsayılan değerlere dön
            profile.background_color = '#F5F5F5'
            profile.text_color = '#000000'
            profile.header_background_color = '#F5F5F5'
            profile.header_text_color = '#333333'
            profile.link_color = '#6E8CA7'
            profile.link_hover_color = '#4E647E'
            profile.button_background_color = '#6E8CA7'
            profile.button_hover_background_color = '#4E647E'
            profile.button_text_color = '#ffffff'
            profile.hover_background_color = '#f0f0f0'
            profile.icon_color = '#333333'
            profile.icon_hover_color = '#007bff'
            profile.answer_background_color = '#F5F5F5'
            profile.content_background_color = '#ffffff'
            profile.tab_background_color = '#f8f9fa'
            profile.tab_text_color = '#000000'
            profile.tab_active_background_color = '#ffffff'
            profile.tab_active_text_color = '#000000'
            profile.dropdown_text_color = '#333333'
            profile.dropdown_hover_background_color = '#f2f2f2'
            profile.dropdown_hover_text_color = '#0056b3'
            profile.nav_link_hover_color = '#007bff'
            profile.nav_link_hover_bg = '#f5f5f5'

            #benim eklediklerim
            profile.message_bubble_color="#d1e7ff"
            profile.tbas_color="#000000"

            profile.font_size='16'

            profile.pagination_background_color = '#ffffff'
            profile.pagination_text_color = '#000000'
            # profile.pagination_active_background_color = '#007bff'
            # profile.pagination_active_text_color = '#ffffff'
            
            profile.cemil = '#ffffff'
            profile.yanit_card='#ffffff' 
            profile.secondary_button_background_color = '#6c757d'
            profile.secondary_button_text_color = '#ffffff'
            profile.secondary_button_hover_background_color= '#495057'

            profile.font_family = 'EB Garamond'
            # Diğer renk alanlarını da varsayılan değerlere ayarlayın
            profile.save()
            messages.success(request, 'Renk ayarlarınız varsayılan değerlere döndürüldü.')
            return redirect('user_settings')
        else:
            # Formdan gelen değerleri kaydet
            profile.secondary_button_background_color = request.POST.get('secondary_button_background_color','#6c757d')
            profile.secondary_button_text_color = request.POST.get('secondary_button_text_color','#ffffff')
            profile.secondary_button_hover_background_color = request.POST.get('secondary_button_hover_background_color','#495057')

            profile.font_size = int(request.POST.get('font_size', 16))  # Sayıya çeviriyoruz

            #benim eklediklerim
            profile.message_bubble_color = request.POST.get('message_bubble_color', '#d1e7ff')
            profile.tbas_color = request.POST.get('tbas_color', '#000000')
            profile.background_color = request.POST.get('background_color', '#F5F5F5')
            profile.text_color = request.POST.get('text_color', '#000000')
            profile.header_background_color = request.POST.get('header_background_color', '#F5F5F5')
            profile.header_text_color = request.POST.get('header_text_color', '#333333')
            profile.link_color = request.POST.get('link_color', '#6E8CA7')
            profile.link_hover_color = request.POST.get('link_hover_color', '#4E647E')
            profile.button_background_color = request.POST.get('button_background_color', '#6E8CA7')
            profile.button_hover_background_color = request.POST.get('button_hover_background_color', '#4E647E')
            profile.button_text_color = request.POST.get('button_text_color', '#ffffff')
            profile.hover_background_color = request.POST.get('hover_background_color', '#f0f0f0')
            profile.icon_color = request.POST.get('icon_color', '#333333')
            profile.icon_hover_color = request.POST.get('icon_hover_color', '#007bff')
            profile.answer_background_color = request.POST.get('answer_background_color', '#F5F5F5')
            profile.content_background_color = request.POST.get('content_background_color', '#ffffff')
            profile.tab_background_color = request.POST.get('tab_background_color', '#f8f9fa')
            profile.tab_text_color = request.POST.get('tab_text_color', '#000000')
            profile.tab_active_background_color = request.POST.get('tab_active_background_color', '#ffffff')
            profile.tab_active_text_color = request.POST.get('tab_active_text_color', '#000000')
            profile.dropdown_text_color = request.POST.get('dropdown_text_color', '#333333')
            profile.dropdown_hover_background_color = request.POST.get('dropdown_hover_background_color', '#f2f2f2')
            profile.dropdown_hover_text_color = request.POST.get('dropdown_hover_text_color', '#0056b3')
            profile.nav_link_hover_color = request.POST.get('nav_link_hover_color', '#007bff')
            profile.nav_link_hover_bg = request.POST.get('nav_link_hover_bg', '#f5f5f5')

            profile.pagination_background_color = request.POST.get('pagination_background_color', '#ffffff')
            profile.pagination_text_color = request.POST.get('pagination_text_color', '#000000')
            # profile.pagination_active_background_color = request.POST.get('pagination_active_background_color', '#007bff')
            # profile.pagination_active_text_color = request.POST.get('pagination_active_text_color', '#ffffff')
            profile.yanit_card= request.POST.get('yanit_card','#ffffff')
            profile.font_family = request.POST.get('font_family', 'EB Garamond')
            # Diğer renk alanlarını da kaydedin
            profile.save()
            messages.success(request, 'Renk ayarlarınız güncellendi.')
            return redirect('user_settings')
    return render(request, 'core/user_settings.html', {'user_profile': profile})

@login_required
def add_starting_question(request):
    all_questions = get_today_questions_queryset()
    if request.method == 'POST':
        form = StartingQuestionForm(request.POST)
        if form.is_valid():
            question_text = form.cleaned_data['question_text']
            question, created = Question.objects.get_or_create(
                question_text=question_text,
                defaults={'user': request.user}
            )
            question.users.add(request.user)
            question.save()
            StartingQuestion.objects.create(user=request.user, question=question)
            Answer.objects.create(
                question=question,
                user=request.user,
                answer_text=form.cleaned_data['answer_text']
            )
            return redirect('user_homepage')
    else:
        form = StartingQuestionForm()
    return render(request, 'core/add_starting_question.html', {'form': form, 'all_questions': all_questions})

@login_required
def vote(request):
    if request.method == 'POST':
        content_type = request.POST.get('content_type')
        object_id = request.POST.get('object_id')
        value = request.POST.get('value')

        if not content_type or not object_id or not value:
            return JsonResponse({'error': 'Missing data'}, status=400)

        try:
            content_type_obj = ContentType.objects.get(model=content_type)
            model_class = content_type_obj.model_class()
            obj = model_class.objects.get(id=object_id)
        except (ContentType.DoesNotExist, model_class.DoesNotExist):
            return JsonResponse({'error': 'Invalid content_type or object_id'}, status=400)

        value = int(value)
        if value not in [1, -1]:
            return JsonResponse({'error': 'Invalid vote value'}, status=400)

        # Get or create the vote
        vote_obj, created = Vote.objects.get_or_create(
            user=request.user,
            content_type=content_type_obj,
            object_id=object_id,
            defaults={'value': value}
        )
        if not created:
            if vote_obj.value == value:
                # Remove the vote if it's the same
                vote_obj.delete()
                value = 0
            else:
                # Update the vote value
                vote_obj.value = value
                vote_obj.save()

        # Recalculate total upvotes and downvotes
        upvotes = Vote.objects.filter(content_type=content_type_obj, object_id=object_id, value=1).count()
        downvotes = Vote.objects.filter(content_type=content_type_obj, object_id=object_id, value=-1).count()

        # Update the object's vote counts
        obj.upvotes = upvotes
        obj.downvotes = downvotes
        obj.save()

        return JsonResponse({
            'upvotes': upvotes,
            'downvotes': downvotes,
            'user_vote_value': value
        })
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)

@login_required
def save_item(request):
    if request.method == 'POST':
        content_type = request.POST.get('content_type')
        object_id = request.POST.get('object_id')

        if not content_type or not object_id:
            return JsonResponse({'error': 'Missing content_type or object_id'}, status=400)

        try:
            content_type_obj = ContentType.objects.get(model=content_type)
        except ContentType.DoesNotExist:
            return JsonResponse({'error': 'Invalid content_type'}, status=400)

        # Daha önce kaydedilmiş mi kontrol edin
        existing_item = SavedItem.objects.filter(
            user=request.user,
            content_type=content_type_obj,
            object_id=object_id
        ).first()

        if existing_item:
            # Eğer kaydedilmişse, kaydı silerek "kaydetmeyi kaldır" işlemi yapın
            existing_item.delete()
            # Kaydedilme sayısını alın
            save_count = SavedItem.objects.filter(
                content_type=content_type_obj,
                object_id=object_id
            ).count()
            return JsonResponse({'status': 'unsaved', 'save_count': save_count})
        else:
            # Yeni bir kayıt oluşturun
            saved_item = SavedItem.objects.create(
                user=request.user,
                content_type=content_type_obj,
                object_id=object_id
            )
            # Kaydedilme sayısını alın
            save_count = SavedItem.objects.filter(
                content_type=content_type_obj,
                object_id=object_id
            ).count()
            return JsonResponse({'status': 'saved', 'save_count': save_count})
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)


def site_statistics(request):
    # Kullanıcı sayısı (en az bir soru veya yanıt yazmış olanlar)
    user_count = User.objects.filter(
        Q(questions__isnull=False) | Q(answers__isnull=False)
    ).distinct().count()

    # Toplam soru ve yanıt sayısı
    total_questions = Question.objects.count()
    total_answers = Answer.objects.count()

    # Toplam beğeni ve beğenmeme sayısı
    total_likes = Vote.objects.filter(value=1).count()
    total_dislikes = Vote.objects.filter(value=-1).count()

    # En çok soru soran kullanıcılar
    top_question_users = User.objects.annotate(
        question_count=Count('questions')
    ).order_by('-question_count')[:5]

    # En çok yanıt veren kullanıcılar
    top_answer_users = User.objects.annotate(
        answer_count=Count('answers')
    ).order_by('-answer_count')[:5]

    # En çok beğenilen sorular (upvotes)
    top_liked_questions = Question.objects.annotate(
        like_count=F('upvotes')
    ).order_by('-like_count')[:5]

    # En çok beğenilen yanıtlar (upvotes)
    top_liked_answers = Answer.objects.annotate(
        like_count=F('upvotes')
    ).order_by('-like_count')[:5]

    # En çok kaydedilen sorular
    top_saved_questions = Question.objects.annotate(
        save_count=Count('saveditem')
    ).order_by('-save_count')[:5]

    # En çok kaydedilen yanıtlar
    top_saved_answers = Answer.objects.annotate(
        save_count=Count('saveditem')
    ).order_by('-save_count')[:5]

    active_tab = request.GET.get('tab', 'word-analysis')  # Varsayılan olarak "Kelime Analizi" sekmesi

    # --- KELİME ANALİZİ ve KELİME ARAMA ---
    question_texts = Question.objects.values_list('question_text', flat=True)
    answer_texts = Answer.objects.values_list('answer_text', flat=True)
    all_texts = list(question_texts) + list(answer_texts)

    # Küçük harfe çevir, birleştir (tüm kelimeleri görmek için)
    all_words = []
    for text in all_texts:
        if text:
            all_words.extend(re.findall(r'\b\w+\b', text.lower()))

    # Hariç tutulan kelimeleri işle
    exclude_words_input = request.GET.get('exclude_words', '')
    if exclude_words_input:
        exclude_words_list = re.split(r',\s*', exclude_words_input.strip())
        exclude_words = set(word.lower() for word in exclude_words_list if word.strip())
    else:
        exclude_words = set()

    # Hariç tutulanları çıkar
    filtered_words = [word for word in all_words if word not in exclude_words]

    # En çok geçen kelimeler
    word_counts = Counter(filtered_words)
    top_words = word_counts.most_common(10)

    # --- Anahtar kelime arama: tüm başlık ve yanıtlarda toplam kaç kere geçiyor? ---
    search_word = request.GET.get('search_word', '').strip().lower()
    search_word_count = None
    if search_word:
        # Hariç tutulan bir kelimeyle aynıysa, doğrudan 0 göster
        if search_word in exclude_words:
            search_word_count = 0
        else:
            # Tek tek başlık ve yanıtların hepsini say (her satırda birden fazla geçiş varsa onlar da dahil)
            search_word_pattern = re.compile(r'\b{}\b'.format(re.escape(search_word)), re.IGNORECASE)
            search_word_count = 0
            for text in all_texts:
                # Hariç tutulan kelimelerden biri bu text'te geçiyorsa, burayı tamamen atla
                if any(re.search(r'\b{}\b'.format(re.escape(exw)), text, re.IGNORECASE) for exw in exclude_words):
                    continue
                if text:
                    search_word_count += len(search_word_pattern.findall(text))

    # --- En çok kullanılan kaynaklar ---

    all_references = list(Reference.objects.all())
    all_references.sort(key=lambda ref: ref.get_usage_count(), reverse=True)
    paginator_references = Paginator(all_references, 5)
    reference_page_number = request.GET.get('reference_page', 1)
    top_references = paginator_references.get_page(reference_page_number)

    context = {
        'active_tab': active_tab,
        'user_count': user_count,
        'total_questions': total_questions,
        'total_answers': total_answers,
        'total_likes': total_likes,
        'total_dislikes': total_dislikes,
        'top_question_users': top_question_users,
        'top_answer_users': top_answer_users,
        'top_liked_questions': top_liked_questions,
        'top_liked_answers': top_liked_answers,
        'top_saved_questions': top_saved_questions,
        'top_saved_answers': top_saved_answers,
        'top_words': top_words,
        'search_word_count': search_word_count,
        'search_word': search_word,
        'exclude_words': ', '.join(sorted(exclude_words)),
        'exclude_words_input': exclude_words_input,
        'top_references': top_references,
    }
    return render(request, 'core/site_statistics.html', context)


def get_today_questions_queryset():
    """
    Son 7 gün içinde oluşturulan veya yanıt alan soruları döndürür (queryset).
    En son yanıtlananlar en üste gelecek şekilde sıralar.
    """
    seven_days_ago = now() - timedelta(days=7)
    queryset = Question.objects.annotate(
        answers_count=Count('answers', distinct=True),  # Yanıt sayısı
        latest_answer_date=Max('answers__created_at')   # En son yanıt tarihi
    ).filter(
        Q(created_at__gte=seven_days_ago) | Q(answers__created_at__gte=seven_days_ago)
    ).distinct()

    queryset = queryset.annotate(
        sort_date=Coalesce('latest_answer_date', 'created_at')
    ).order_by(F('sort_date').desc())

    return queryset

def get_today_questions_page(request, per_page=25):
    queryset = get_today_questions_queryset()
    paginator = Paginator(queryset, per_page)
    page = request.GET.get('page', 1)
    try:
        all_questions = paginator.page(page)
    except PageNotAnInteger:
        all_questions = paginator.page(1)
    except EmptyPage:
        all_questions = paginator.page(paginator.num_pages)
    return all_questions

def user_homepage(request):
    if not request.user.is_authenticated:
        return redirect('signup')

    # 1. Tüm Sorular (ve cevap sayısı)
    all_questions_qs = get_today_questions_queryset().select_related('user')
    q_page_number = request.GET.get('page', 1)
    q_paginator = Paginator(all_questions_qs, 20)
    all_questions = q_paginator.get_page(q_page_number)

    # 2. Rastgele Cevaplar -- OPTİMİZE (ID ile çek, sonra select_related)
    answer_ids = list(Answer.objects.values_list('id', flat=True))
    if len(answer_ids) > 20:
        random_ids = random.sample(answer_ids, 20)
    else:
        random_ids = answer_ids
    random_items_qs = Answer.objects.filter(id__in=random_ids).select_related('question', 'user')

    # Sıralama tutarsız olmasın diye aynı sırayı koru:
    random_items = sorted(list(random_items_qs), key=lambda x: random_ids.index(x.id)) if random_ids else []

    # 3. Başlangıç Soruları (her biri için toplam subquestion ve son eklenen alt soru)
    starting_questions_qs = StartingQuestion.objects.filter(user=request.user).select_related('question').annotate(
        total_subquestions=Count('question__subquestions'),
        latest_subquestion_date=Max('question__subquestions__created_at')
    ).order_by(F('latest_subquestion_date').desc(nulls_last=True)).select_related('question__user')

    sq_page_number = request.GET.get('starting_page', 1)
    sq_paginator = Paginator(starting_questions_qs, 10)
    starting_questions = sq_paginator.get_page(sq_page_number)

    # 4. Kullanıcı oyları topluca çek
    content_type_answer = ContentType.objects.get_for_model(Answer)
    if request.user.is_authenticated and random_items:
        answer_votes = Vote.objects.filter(
            user=request.user,
            content_type=content_type_answer,
            object_id__in=[a.id for a in random_items]
        ).values('object_id', 'value')
        answer_vote_dict = {item['object_id']: item['value'] for item in answer_votes}
        for answer in random_items:
            answer.user_vote_value = answer_vote_dict.get(answer.id, 0)
    else:
        for answer in random_items:
            answer.user_vote_value = 0

    # 5. Kullanıcının kaydettiği yanıtlar
    if request.user.is_authenticated and random_items:
        saved_answer_ids = set(SavedItem.objects.filter(
            user=request.user,
            content_type=content_type_answer,
            object_id__in=[a.id for a in random_items]
        ).values_list('object_id', flat=True))
    else:
        saved_answer_ids = set()

    # 6. Yanıtların kaydedilme sayısı
    answer_save_counts = SavedItem.objects.filter(
        content_type=content_type_answer,
        object_id__in=[a.id for a in random_items]
    ).values('object_id').annotate(count=Count('id'))
    answer_save_dict = {item['object_id']: item['count'] for item in answer_save_counts}

    context = {
        'random_items': random_items,
        'saved_answer_ids': saved_answer_ids,
        'answer_save_dict': answer_save_dict,
        'all_questions': all_questions,  # Pagineli queryset
        'starting_questions': starting_questions,  # Pagineli queryset
    }
    return render(request, 'core/user_homepage.html', context)

@login_required
def edit_answer(request, answer_id):
    all_questions = get_today_questions_queryset()
    answer = get_object_or_404(Answer, id=answer_id, user=request.user)
        # Başlangıç sorularını al
    starting_questions = StartingQuestion.objects.filter(user=request.user).annotate(
        total_subquestions=Count('question__subquestions'),
        latest_subquestion_date=Max('question__subquestions__created_at')
    ).order_by(F('latest_subquestion_date').desc(nulls_last=True))

    if request.method == 'POST':
        form = AnswerForm(request.POST, instance=answer)
        if form.is_valid():
            form.save()
            messages.success(request, 'Yanıt başarıyla güncellendi.')
            return redirect('question_detail', question_id=answer.question.id)
    else:
        form = AnswerForm(instance=answer)

    return render(request, 'core/edit_answer.html', {'form': form, 'answer': answer,'all_questions': all_questions,'starting_questions': starting_questions,})


@login_required
def delete_question(request, question_id):
    question = get_object_or_404(Question, id=question_id)

    if request.method == 'POST':
        if request.user == question.user:
            with transaction.atomic():
                # Delete all answers associated with the question by the user
                Answer.objects.filter(question=question, user=request.user).delete()
                # Remove the user from the question's users
                question.users.remove(request.user)
                if question.users.count() == 0:
                    # If no users are associated, delete the question and its subquestions
                    delete_question_and_subquestions(question)
                    messages.success(request, 'Soru ve alt soruları başarıyla silindi.')
                else:
                    messages.success(request, 'Soru sizin için silindi.')
            return redirect('user_homepage')
        else:
            messages.error(request, 'Bu soruyu silme yetkiniz yok.')
            return redirect('question_detail', question_id=question.id)
    else:
        return render(request, 'core/confirm_delete_question.html', {'question': question})

@login_required
def delete_answer(request, answer_id):
    answer = get_object_or_404(Answer, id=answer_id, user=request.user)
    next_url = request.GET.get('next', '')  # ?next=...
    
    if request.method == 'POST':
        answer.delete()
        if next_url:
            return redirect(unquote(next_url))
        else:
            # Varsayılan davranış (örneğin ana sayfa):
            return redirect('user_homepage')
    else:
        if next_url:
            return redirect(unquote(next_url))
        else:
            return redirect('user_homepage')

def delete_question_and_subquestions(question):
    subquestions = question.subquestions.all()
    for sub in subquestions:
        delete_question_and_subquestions(sub)
    question.delete()

@login_required
def single_answer(request, question_id, answer_id):
    question = get_object_or_404(Question, id=question_id)
    focused_answer = get_object_or_404(Answer, id=answer_id, question=question)
    all_questions = get_today_questions_queryset()

    # Tüm yanıtlar
    all_answers = Answer.objects.filter(question=question).select_related('user')

    # Kullanıcının kaydettiği yanıtlar
    saved_answer_ids = SavedItem.objects.filter(
        user=request.user,
        content_type=ContentType.objects.get_for_model(Answer),
        object_id__in=all_answers.values_list('id', flat=True)
    ).values_list('object_id', flat=True)

    # Yanıtların kaydedilme sayısı
    content_type_answer = ContentType.objects.get_for_model(Answer)
    answer_save_counts = SavedItem.objects.filter(
        content_type=content_type_answer,
        object_id__in=all_answers.values_list('id', flat=True)
    ).values('object_id').annotate(count=Count('id'))
    answer_save_dict = {item['object_id']: item['count'] for item in answer_save_counts}

    # Kullanıcının oyları
    user_votes = Vote.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=all_answers.values_list('id', flat=True)
    ).values('object_id', 'value')
    user_vote_dict = {vote['object_id']: vote['value'] for vote in user_votes}

    # Oy değerlerini her yanıta ekle
    for ans in all_answers:
        ans.user_vote_value = user_vote_dict.get(ans.id, 0)

    # Yanıt ekleme formu
    if request.method == "POST":
        form = AnswerForm(request.POST)
        if form.is_valid():
            new_answer = form.save(commit=False)
            new_answer.question = question
            new_answer.user = request.user
            new_answer.save()
            return redirect('single_answer', question_id=question.id, answer_id=new_answer.id)
    else:
        form = AnswerForm()

    context = {
        'question': question,
        'focused_answer': focused_answer,
        'all_answers': all_answers,
        'saved_answer_ids': saved_answer_ids,
        'answer_save_dict': answer_save_dict,
        'form': form,  # Yanıt ekleme formu
        'all_questions': all_questions,
    }
    return render(request, 'core/single_answer.html', context)

def user_search(request):
    query = request.GET.get('q', '').strip()
    users = User.objects.filter(username__icontains=query)[:10]
    results = [{'id': user.id, 'username': user.username} for user in users]
    return JsonResponse({'results': results})


def map_data_view(request):
    user_ids = request.GET.getlist('user_id')
    filter_param = request.GET.get('filter')

    # Sadece başlangıç veya alt soru olanlar
    starting_question_ids = StartingQuestion.objects.values_list('question_id', flat=True)
    subquestion_ids = Question.objects.filter(parent_questions__isnull=False).values_list('id', flat=True)
    question_ids = set(starting_question_ids) | set(subquestion_ids)

    # Filtre uygula
    if filter_param == 'me':
        questions = Question.objects.filter(id__in=question_ids, user=request.user)
    elif user_ids:
        questions = Question.objects.filter(id__in=question_ids, user__id__in=user_ids).distinct()
    else:
        questions = Question.objects.filter(id__in=question_ids)

    # Dönüş yap
    data = generate_question_nodes(questions)
    return JsonResponse(data, safe=False)

def generate_question_nodes(questions):
    

    nodes = {}
    links = []
    question_text_to_ids = defaultdict(list)

    # Düğümleri oluştur
    for question in questions:
        key = question.question_text
        question_text_to_ids[key].append(question.id)
        if key not in nodes:
            # Kullanıcıların yanıtlarını da getir
            user_entries = []
            # Bu soruya ait tüm yanıtları çek (user, answer_id)
            answers = question.answers.all().select_related('user')
            for answer in answers:
                user_entries.append({
                    "id": answer.user.id,
                    "username": answer.user.username,
                    "answer_id": answer.id,
                })

            user_ids = [entry["id"] for entry in user_entries]
            node = {
                "id": f"q{hash(key)}",
                "label": question.question_text,
                "users": user_entries,  # Burada artık her kullanıcı için answer_id de var
                "size": 20 + 10 * (len(user_entries) - 1),
                "color": '',
                "question_id": question.id,
                "question_ids": [question.id],
            }
            # Renk belirleme mantığı
            if len(user_ids) == 1:
                node["color"] = get_user_color(user_ids[0])
            elif len(user_ids) > 1:
                node["color"] = '#CCCCCC'
            else:
                node["color"] = '#000000'
            nodes[key] = node
        else:
            # Kullanıcıları ve boyutu güncelle
            existing_node = nodes[key]
            # Mevcut ve yeni user_entries birleştirilir, answer_id ile benzersizleştirilebilir
            new_user_entries = []
            answers = question.answers.all().select_related('user')
            for answer in answers:
                new_user_entries.append({
                    "id": answer.user.id,
                    "username": answer.user.username,
                    "answer_id": answer.id,
                })
            # Duplicate user id'leri önlemek için
            user_dict = {entry["id"]: entry for entry in (existing_node["users"] + new_user_entries)}
            combined_user_entries = list(user_dict.values())
            existing_node["users"] = combined_user_entries
            existing_node["size"] = 20 + 5 * (len(combined_user_entries) - 1)
            existing_node["question_ids"].append(question.id)
            # Renk güncelle
            if len(combined_user_entries) == 1:
                existing_node["color"] = get_user_color(combined_user_entries[0]["id"])
            elif len(combined_user_entries) > 1:
                existing_node["color"] = '#CCCCCC'
            else:
                existing_node["color"] = '#000000'

    # Bağlantılar
    link_set = set()
    for question in questions:
        source_key = question.question_text
        for subquestion in question.subquestions.all():
            target_key = subquestion.question_text
            if target_key in nodes:
                link_id = (nodes[source_key]["id"], nodes[target_key]["id"])
                if link_id not in link_set:
                    links.append({
                        "source": nodes[source_key]["id"],
                        "target": nodes[target_key]["id"]
                    })
                    link_set.add(link_id)

    question_nodes = {
        "nodes": list(nodes.values()),
        "links": links
    }
    return question_nodes


@login_required
def bkz_view(request, query):
    try:
        question = Question.objects.get(question_text__iexact=query)
        return redirect('question_detail', question_id=question.id)
    except Question.DoesNotExist:
        # Soru bulunamazsa, add_question_from_search sayfasına yönlendirin
        return redirect(f'{reverse("add_question_from_search")}?q={query}')

@login_required
def reference_search(request):
    query = request.GET.get('q', '').strip()
    results = []
    if query:
        questions = Question.objects.filter(question_text__icontains=query)
        for question in questions:
            results.append({
                'id': question.id,
                'text': question.question_text
            })
    return JsonResponse({'results': results})

@login_required
def pin_entry(request, answer_id):
    if request.method == 'POST':
        user = request.user
        # Mevcut sabitlenmiş girdiyi kaldır
        PinnedEntry.objects.filter(user=user).delete()
        # Yeni girdiyi sabitle
        answer = get_object_or_404(Answer, id=answer_id)
        PinnedEntry.objects.create(user=user, answer=answer)
    return redirect(request.META.get('HTTP_REFERER', 'home'))

@login_required
def unpin_entry(request):
    if request.method == 'POST':
        user = request.user
        PinnedEntry.objects.filter(user=user).delete()
    return redirect(request.META.get('HTTP_REFERER', 'home'))

@login_required
def update_profile_photo(request):
    profile_user = request.user
    user_profile = profile_user.userprofile

    if request.method == 'POST':
        form = ProfilePhotoForm(request.POST, request.FILES, instance=user_profile)
        remove_photo = (request.POST.get('remove_photo') == 'true')

        if remove_photo:
            if user_profile.photo:
                user_profile.photo.delete(save=False)
            user_profile.photo = None
            user_profile.save()
            messages.success(request, 'Fotoğrafınız kaldırıldı.')
            return redirect('user_profile', username=profile_user.username)

        if form.is_valid():
            form.save()
            messages.success(request, 'Profil fotoğrafınız güncellendi.')
            return redirect('user_profile', username=profile_user.username)
        else:
            # Form hatalıysa modal tekrar açılsın diye user_profile.html'i tekrar render ediyoruz.
            return render(request, 'core/user_profile.html', {
                'profile_user': profile_user,
                'user_profile': user_profile,
                'form': form,
                'is_own_profile': True,  # Kendi profilimiz olduğunu varsayıyoruz
            })
    else:
        # GET isteğinde direkt profiline yönlendir
        return redirect('user_profile', username=profile_user.username)


def get_top_words(user):
    answers = Answer.objects.filter(user=user)
    questions = Question.objects.filter(user=user)

    text = ' '.join([a.answer_text for a in answers] + [q.question_text for q in questions])
    words = re.findall(r'\w+', text.lower())
    word_counts = Counter(words)
    top_words = word_counts.most_common(10)

    return top_words


@login_required
@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def message_list(request):
    messages = Message.objects.filter(
        Q(sender=request.user) | Q(recipient=request.user)
    ).select_related('sender', 'recipient')

    # En yeni mesajı her kullanıcı için birleştir
    # Sadece diğer kullanıcı id’lerini çek
    user_ids = set(messages.values_list('sender', flat=True)) | set(messages.values_list('recipient', flat=True))
    user_ids.discard(request.user.id)
    other_users = User.objects.filter(id__in=user_ids)

    conversation_dict = {}
    for other_user in other_users:
        messages_with_user = messages.filter(
            Q(sender=other_user, recipient=request.user) |
            Q(sender=request.user, recipient=other_user)
        ).order_by('-timestamp')

        unread_count = messages_with_user.filter(
            sender=other_user,
            recipient=request.user,
            is_read=False
        ).count()

        conversation_dict[other_user] = {
            'messages': messages_with_user,
            'unread_count': unread_count,
        }

    all_questions = get_today_questions_queryset().annotate(
        answers_count=Count('answers')
    )

    context = {
        'conversations': conversation_dict,
        'all_questions': all_questions,
    }
    return render(request, 'core/message_list.html', context)

@login_required
def message_detail(request, username):
    other_user = get_object_or_404(User, username=username)
    Message.objects.filter(sender=other_user, recipient=request.user, is_read=False).update(is_read=True)

    # Mesajları doğru şekilde sıralayın: en eski önce, en yeni sonra
    messages = Message.objects.filter(
        Q(sender=request.user, recipient=other_user) |
        Q(sender=other_user, recipient=request.user)
    ).order_by('timestamp')  # 'timestamp' alanına göre artan sırada
    
    if request.method == 'POST':
        body = request.POST.get('body')
        if body:
            Message.objects.create(
                sender=request.user,
                recipient=other_user,
                body=body,
                timestamp=timezone.now(),
                is_read=False
            )
            # Diğer kullanıcının mesajlarını okunmuş olarak işaretlemek isteğe bağlıdır
            Message.objects.filter(sender=other_user, recipient=request.user).update(is_read=True)
            return redirect('message_detail', username=username)
    all_questions = get_today_questions_queryset()

    context = {
        'other_user': other_user,
        'messages': messages,
        'all_questions': all_questions,
    }
    return render(request, 'core/message_detail.html', context)

@login_required
def send_message_from_answer(request, answer_id):
    answer = get_object_or_404(Answer, id=answer_id)
    recipient = answer.user

    # İlgili yanıta ait path ve tam URL
    
    answer_url_path = reverse('single_answer', args=[answer.question.id, answer.id])
    answer_full_url = request.build_absolute_uri(answer_url_path)

    if request.method == 'POST':
        form = MessageForm(request.POST)
        if form.is_valid():
            message = form.save(commit=False)
            message.sender = request.user
            message.recipient = recipient
            # Sadece POST aşamasında, bir kere ekliyoruz:
            message.body = f"{answer_full_url} {message.body}"
            message.timestamp = timezone.now()
            message.is_read = False
            message.save()
            return redirect('message_detail', username=recipient.username)
    else:
        # GET aşamasında link eklemiyoruz, sadece boş bir form dönüyoruz.
        form = MessageForm(initial={'recipient': recipient})

    context = {
        'form': form,
        'recipient': recipient,
        'answer': answer,
    }
    return render(request, 'core/send_message_from_answer.html', context)

@login_required
def check_new_messages(request):
    # Kullanıcının okunmamış mesajlarını say
    unread_count = Message.objects.filter(recipient=request.user, is_read=False).count()
    return JsonResponse({'unread_count': unread_count})

@login_required
def send_message_from_user(request, user_id):
    recipient = get_object_or_404(User, id=user_id)

    if request.method == 'POST':
        form = MessageForm(request.POST)
        if form.is_valid():
            message = form.save(commit=False)
            message.sender = request.user
            message.recipient = recipient
            message.save()
            return redirect('message_detail', username=recipient.username)
    else:
        form = MessageForm()

    context = {
        'form': form,
        'recipient': recipient,
    }
    return render(request, 'core/send_message_from_answer.html', context)


def custom_404_view(request, exception):
    # 404 sayfası için özel view
    # Status kodunu 404 olarak ayarlamayı unutmayın.
    response = render(request, 'core/404.html')
    response.status_code = 404
    return response

@require_POST
@login_required
def ignore_random_sentence(request):
    """
    POST ile gelen sentence_id parametresi üzerinden
    cümleyi ignore listesine ekler.
    """
    sentence_id = request.POST.get('sentence_id')
    if not sentence_id:
        return JsonResponse({'status': 'error', 'message': 'sentence_id parametresi bulunamadı.'}, status=400)

    try:
        sentence_obj = RandomSentence.objects.get(id=sentence_id)
    except RandomSentence.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Cümle bulunamadı.'}, status=404)

    # Kullanıcıyı ignore listesine ekle
    sentence_obj.ignored_by.add(request.user)
    return JsonResponse({'status': 'success', 'message': 'Cümle ignore listesine eklendi.'})

def get_random_sentence(request):
    if request.user.is_authenticated:
        # ignore ettiği cümleler hariç
        sentences = RandomSentence.objects.exclude(ignored_by=request.user)
    else:
        # anonimse hepsi arasından rastgele çek
        sentences = RandomSentence.objects.all()

    sentences_count = sentences.count()
    if sentences_count == 0:
        return JsonResponse({'sentence': 'Henüz eklenmiş (veya gösterilmeye uygun) cümle yok.'})

    random_index = random.randint(0, sentences_count - 1)
    sentence_obj = sentences[random_index]

    return JsonResponse({
        'id': sentence_obj.id,
        'sentence': sentence_obj.sentence
    })

@require_POST
def add_random_sentence(request):
    form = RandomSentenceForm(request.POST)
    if form.is_valid():
        form.save()
        return JsonResponse({'status': 'success', 'message': 'Cümle eklendi!'})
    else:
        # Form geçersiz ise hata mesajını döndür
        errors = form.errors.get_json_data()
        return JsonResponse({'status': 'error', 'errors': errors})
    
#POLLS
@login_required
def polls_home(request):
    # Aktif anketler (end_date > şu an)
    active_polls = Poll.objects.filter(end_date__gt=timezone.now()).order_by('-created_at')
    # Süresi geçmiş anketler (end_date <= şu an)
    expired_polls = Poll.objects.filter(end_date__lte=timezone.now()).order_by('-created_at')

    # Süresi geçmiş anketler için yüzdeleri hesaplayalım
    expired_polls_data = []
    for poll in expired_polls:
        options_data = []
        total_votes = sum([opt.votes.count() for opt in poll.options.all()])
        for opt in poll.options.all():
            if total_votes > 0:
                percentage = (opt.votes.count() / total_votes) * 100
            else:
                percentage = 0
            options_data.append({
                'text': opt.option_text,
                'percentage': f"{percentage:.2f}"
            })
        expired_polls_data.append({
            'poll': poll,
            'options_data': options_data
        })

    form = PollForm()
    return render(request, 'core/polls.html', {
        'active_polls': active_polls,
        'expired_polls_data': expired_polls_data,
        'form': form
    })

@login_required
def create_poll(request):
    if request.method == 'POST':
        form = PollForm(request.POST)
        if form.is_valid():
            question_text = form.cleaned_data['question_text']
            end_date = form.cleaned_data['end_date']
            is_anonymous = form.cleaned_data['is_anonymous']
            options = form.cleaned_data['options']

            poll = Poll.objects.create(
                question_text=question_text,
                created_by=request.user,
                end_date=end_date,
                is_anonymous=is_anonymous
            )
            for opt_text in options:
                PollOption.objects.create(poll=poll, option_text=opt_text)
            
            messages.success(request, 'Anket başarıyla oluşturuldu.')
            return redirect('polls_home')
        else:
            # Form geçersizse hata mesajlarını göster
            return render(request, 'core/create_poll.html', {'form': form})
    else:
        form = PollForm()
        return render(request, 'core/create_poll.html', {'form': form})

@login_required
def vote_poll(request, poll_id, option_id):
    poll = get_object_or_404(Poll, id=poll_id)
    option = get_object_or_404(PollOption, id=option_id, poll=poll)

    if PollVote.objects.filter(user=request.user, option__poll=poll).exists():
        messages.error(request, 'Bu ankete daha önce oy verdiniz.')
        return redirect('polls_home')

    # -------- Tarih kontrolü --------
    if poll.end_date and timezone.now() > poll.end_date:
        messages.error(request, 'Bu anketin süresi dolmuş.')
        return redirect('polls_home')

    if poll.is_active():
        PollVote.objects.create(user=request.user, option=option)
        messages.success(request, 'Oyunuz kaydedildi.')
    else:
        messages.error(request, 'Bu anket süresi dolmuş.')
    return redirect('polls_home')


@login_required
def poll_question_redirect(request, poll_id):
    poll = get_object_or_404(Poll, id=poll_id)
    if poll.related_question:
        # Soru mevcutsa direkt oraya git
        return redirect('question_detail', question_id=poll.related_question.id)
    else:
        # Yeni başlık oluştur
        q = Question.objects.create(
            question_text=f"anket:{poll.question_text}",
            user=request.user
        )
        q.users.add(request.user)
        poll.related_question = q
        poll.save()
        return redirect('question_detail', question_id=q.id)
    
@login_required
def vote_poll_ajax(request, poll_id):
    poll = get_object_or_404(Poll, id=poll_id)

    # ---- Tarih kontrolü EKLENECEK! ----
    if poll.end_date and timezone.now() > poll.end_date:
        # Süresi geçmiş: Sonucu döndür + hata mesajı
        options = poll.options.all()
        total_votes = sum(opt.votes.count() for opt in options)
        user_vote = PollVote.objects.filter(option__poll=poll, user=request.user).first()
        context = {
            'poll': poll,
            'options': options,
            'total_votes': total_votes,
            'user_vote': user_vote,
            'poll_expired': True,  # yeni context anahtarı!
        }
        html = render_to_string('core/poll_popover_content.html', context, request)
        return JsonResponse({'html': html, 'error': "Bu anketin süresi doldu."}, status=400)

    # --- Normal oy verme ---
    if request.method == 'POST':
        option_id = request.POST.get('option_id')
        if PollVote.objects.filter(user=request.user, option__poll=poll).exists():
            # Zaten oy vermiş
            pass
        else:
            option = get_object_or_404(PollOption, id=option_id, poll=poll)
            PollVote.objects.create(user=request.user, option=option)
    # Sonuçları tekrar gönder
    options = poll.options.all()
    total_votes = sum(opt.votes.count() for opt in options)
    user_vote = PollVote.objects.filter(option__poll=poll, user=request.user).first()
    context = {
        'poll': poll,
        'options': options,
        'total_votes': total_votes,
        'user_vote': user_vote,
        'poll_expired': False,  # yeni context anahtarı!
    }
    html = render_to_string('core/poll_popover_content.html', context, request)
    return JsonResponse({'html': html})

@login_required
def poll_detail(request, poll_id):
    poll = get_object_or_404(Poll, id=poll_id)
    options = poll.options.all()
    total_votes = sum(opt.votes.count() for opt in options)
    user_vote = None

    if request.user.is_authenticated:
        # Kullanıcı daha önce oy vermiş mi?
        for opt in options:
            if opt.votes.filter(user=request.user).exists():
                user_vote = opt.id
                break

    # Oy kullanma işlemi (POST)
    if request.method == "POST":
        if not user_vote:  # Sadece ilk oy
            option_id = request.POST.get("option")
            option = get_object_or_404(PollOption, id=option_id, poll=poll)
            PollVote.objects.create(user=request.user, option=option)
            return redirect('poll_detail', poll_id=poll.id)
        else:
            messages.error(request, "Bu ankete zaten oy verdiniz.")
            return redirect('poll_detail', poll_id=poll.id)

    # Her seçenek için yüzdeleri hesapla
    options_data = []
    for opt in options:
        votes = opt.votes.count()
        percentage = (votes / total_votes * 100) if total_votes > 0 else 0
        options_data.append({
            'option': opt,
            'votes': votes,
            'percentage': round(percentage, 2)
        })

    context = {
        'poll': poll,
        'options_data': options_data,
        'total_votes': total_votes,
        'user_vote': user_vote,
    }
    return render(request, 'core/poll_detail.html', context)

@login_required
def poll_popover_content(request, poll_id):
    poll = get_object_or_404(Poll, id=poll_id)
    options = poll.options.all()
    total_votes = sum(opt.votes.count() for opt in options)
    user_vote = None
    if request.user.is_authenticated:
        user_vote = PollVote.objects.filter(user=request.user, option__poll=poll).first()
    # Her option için yüzdeyi hesapla ve diziye ekle:
    options_data = []
    for opt in options:
        votes = opt.votes.count()
        percentage = (votes / total_votes * 100) if total_votes > 0 else 0
        options_data.append({
            'option': opt,
            'votes': votes,
            'percentage': round(percentage, 1)
        })
    context = {
        'poll': poll,
        'options_data': options_data,
        'user_vote': user_vote,
        'poll_expired': poll.end_date and timezone.now() > poll.end_date,
        'total_votes': total_votes,
    }
    html = render_to_string('core/poll_popover_content.html', context)
    return JsonResponse({'html': html})

def create_definition(request, question_id):
    question = get_object_or_404(Question, id=question_id)
    if request.method == 'POST':
        body = request.body.decode('utf-8')
        data = json.loads(body)
        form = DefinitionForm(data)
        if form.is_valid():
            definition_obj = form.save(commit=False)
            definition_obj.user = request.user
            definition_obj.question = question
            definition_obj.save()

            # TANIMI AYNI ZAMANDA ANSWER OLARAK DA KAYDET
            answer_obj = Answer.objects.create(
                question=question,
                user=request.user,
                answer_text=definition_obj.definition_text
            )
            # Oluşturulan answer'ı tanıma bağlayın
            definition_obj.answer = answer_obj
            definition_obj.save()

            return JsonResponse({
                'status': 'success',
                'definition_text': definition_obj.definition_text,
                'question_text': question.question_text
            }, status=200)
        else:
            return JsonResponse({'status': 'error', 'errors': form.errors}, status=400)
    return JsonResponse({'status': 'invalid_method'}, status=405)


@login_required
def get_user_definitions(request):
    """
    Kullanıcının tanımlarını JSON olarak döndürür.
    Arama + sayfalama + usage_count eklenmiştir.
    ?q=abc  => Arama
    ?page=2 => Sayfa numarası
    """
    if request.method == 'GET':
        q = request.GET.get('q', '').strip()
        page_num = request.GET.get('page', '1')

        # Sadece kullanıcıya ait Definition’lar
        defs = Definition.objects.filter(user=request.user)

        # Arama (soru metninde veya tanım metninde?)
        # d.question.question_text veya d.definition_text
        if q:
            defs = defs.filter(
                Q(definition_text__icontains=q) |
                Q(question__question_text__icontains=q)
            )

        # Alfabetik sıralama => question_text üzerinden
        defs = defs.order_by('question__question_text')

        # Sayfalama
        paginator = Paginator(defs, 5)  # sayfa başına 5 kayıt
        try:
            page_obj = paginator.page(page_num)
        except:
            page_obj = paginator.page(1)

        data_list = []
        for d in page_obj.object_list:
            usage_count_self = Answer.objects.filter(
                user=request.user,
                answer_text__icontains=f'(tanim:{d.question.question_text}:{d.id})'
            ).count()
            usage_count_all = Answer.objects.filter(
                answer_text__icontains=f'(tanim:{d.question.question_text}:{d.id})'
            ).count()

            data_list.append({
                'id': d.id,
                'question_id': d.question.id,
                'question_text': d.question.question_text,
                'definition_text': d.definition_text[:80] + '...' if len(d.definition_text) > 80 else d.definition_text,
                'usage_count_self': usage_count_self,  # Bu tanımı bu kullanıcı kaç kere kullanmış
                'usage_count_all': usage_count_all,    # Bu tanımı tüm kullanıcılar kaç kere kullanmış
            })

        response_data = {
            'definitions': data_list,
            'current_page': page_obj.number,
            'total_pages': paginator.num_pages,
        }
        return JsonResponse(response_data, status=200)
    else:
        return JsonResponse({'error': 'invalid method'}, status=405)
@login_required
def edit_definition(request, definition_id):
    definition = get_object_or_404(Definition, id=definition_id, user=request.user)
    if request.method == 'POST':
        # İçerik JSON veya form-data olabilir:
        if request.headers.get('Content-Type') == 'application/json':
            data = json.loads(request.body.decode('utf-8'))
        else:
            data = request.POST
        form = DefinitionForm(data, instance=definition)
        if form.is_valid():
            form.save()
            # Tanı nesnesini DB'den tazeleyelim
            definition.refresh_from_db()
            if definition.answer:
                definition.answer.answer_text = definition.definition_text
                definition.answer.save()
            else:
                # Eğer tanıya bağlı Answer yoksa yeni oluştur
                new_answer = Answer.objects.create(
                    question=definition.question,
                    user=definition.user,
                    answer_text=definition.definition_text
                )
                definition.answer = new_answer
                definition.save()
            messages.success(request, 'Tanım güncellendi.')
            # Değişikliklerin hemen görünmesi için, düzenlemeden sonra ilgili soru detayına yönlendirebilirsiniz:
            return redirect('question_detail', question_id=definition.question.id)
        else:
            messages.error(request, 'Form hataları: %s' % form.errors)
    else:
        form = DefinitionForm(instance=definition)
    
    return render(request, 'core/edit_definition.html', {
        'form': form,
        'definition': definition,
    })

@login_required
def delete_definition(request, definition_id):
    definition = get_object_or_404(Definition, id=definition_id, user=request.user)
    if request.method == 'POST':
        definition.delete()
        messages.success(request, 'Tanım silindi.')
        return redirect(f"{reverse('user_profile', args=[request.user.username])}?tab=tanimlar")
    # “GET” istek geldiğinde doğrulama penceresi (confirm) gösterebilirsiniz.
    return render(request, 'core/confirm_delete_definition.html', {'definition': definition})

@login_required
def get_all_definitions(request):
    """
    Tüm kullanıcıların tanımlarını JSON döndürür.
    ?q= => arama
    ?page= => sayfa
    Aynı şekilde usage_count içerir.
    """
    if request.method == 'GET':
        q = request.GET.get('q', '').strip()
        page_num = request.GET.get('page', '1')

        defs = Definition.objects.select_related('question', 'user').all()

        if q:
            defs = defs.filter(
                Q(definition_text__icontains=q) |
                Q(question__question_text__icontains=q) |
                Q(user__username__icontains=q)
            )

        # Alfabetik sıralama
        defs = defs.order_by('question__question_text')

        # Sayfalama
        paginator = Paginator(defs, 5)  # 5’erli
        try:
            page_obj = paginator.page(page_num)
        except:
            page_obj = paginator.page(1)

        data_list = []
        for d in page_obj.object_list:
            usage_count_all = Answer.objects.filter(
                answer_text__icontains=f'(tanim:{d.question.question_text}:{d.id})'
            ).count()
            data_list.append({
                'id': d.id,
                'question_text': d.question.question_text,
                'definition_text': d.definition_text[:80] + '...' if len(d.definition_text) > 80 else d.definition_text,
                'username': d.user.username,
                'usage_count_all': usage_count_all,
            })

        return JsonResponse({
            'definitions': data_list,
            'current_page': page_obj.number,
            'total_pages': paginator.num_pages,
        }, status=200)
    else:
        return JsonResponse({'error': 'invalid method'}, status=405)

@require_POST
@login_required
def create_reference(request):
    """
    AJAX ile yeni bir Reference (Kaynak) oluşturmak için.
    """
    form = ReferenceForm(request.POST)
    if form.is_valid():
        # commit=False ile kaynağı oluşturuyoruz, ardından created_by alanını ekliyoruz
        reference_obj = form.save(commit=False)
        reference_obj.created_by = request.user
        reference_obj.save()
        data = {
            'status': 'success',
            'reference': {
                'id': reference_obj.id,
                'display': str(reference_obj),
            }
        }
        return JsonResponse(data, status=200)
    else:
        return JsonResponse({'status': 'error', 'errors': form.errors}, status=400)

def get_references(request):
    """
    Tüm referans kayıtlarını, opsiyonel 'q' arama parametresi ile 
    filtreleyerek JSON döndürür.
    """
    q = request.GET.get('q', '').strip()  # Arama terimi
    references = Reference.objects.all()

    if q:
        references = references.filter(
            Q(author_surname__icontains=q) |
            Q(author_name__icontains=q) |
            Q(rest__icontains=q) |
            Q(abbreviation__icontains=q)|
            Q(metin_ismi__icontains=q)|
            Q(year__iexact=q)
            # Yıl alanında arama yapmak isterseniz (IntegerField olduğu için tam eşleşmede kullanabilirsiniz)
            # veya q yalnızca sayıysa year__icontains gibi bir yaklaşımla handle edebilirsiniz.
            # Örnek:
            # 
        )
    
    # Sıralama
    references = references.order_by('author_surname', 'year')

    data = []
    for ref in references:
        data.append({
            'id': ref.id,
            'author_surname': ref.author_surname,
            'author_name': ref.author_name,
            'year': ref.year,
            'rest': ref.rest,
            'abbreviation': ref.abbreviation or '',
            'display': str(ref),
        })
    
    return JsonResponse({'references': data}, status=200)

@login_required
def edit_reference(request, reference_id):
    # Yalnızca kaynağı oluşturan kullanıcı düzenleyebilsin
    reference = get_object_or_404(Reference, id=reference_id, created_by=request.user)
    if request.method == 'POST':
        form = ReferenceForm(request.POST, instance=reference)
        if form.is_valid():
            form.save()
            messages.success(request, "Kaynak başarıyla güncellendi.")
            return redirect('user_profile', username=request.user.username)
    else:
        form = ReferenceForm(instance=reference)
    return render(request, 'core/edit_reference.html', {'form': form})

@login_required
def delete_reference(request, reference_id):
    # Yalnızca kaynağı oluşturan kullanıcı silebilsin
    reference = get_object_or_404(Reference, id=reference_id, created_by=request.user)
    if request.method == 'POST':
        reference.delete()
        messages.success(request, "Kaynak başarıyla silindi.")
        return redirect('user_profile', username=request.user.username)
    return render(request, 'core/confirm_delete_reference.html', {'reference': reference})

def download_entries_json(request, username):
    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse(
            {'error': 'Bu işlemi yapmaya yetkiniz yok.'},
            status=403
        )

    user_questions = Question.objects.filter(user=target_user).order_by('created_at')

    questions_data = []
    for q in user_questions:
        # Soru sahibinin kendi yanıtlarını çekmek istiyorsanız:
        q_answers = q.answers.filter(user=target_user).order_by('created_at')
        # Tüm yanıtları (herhangi bir kullanıcıdan) istiyorsanız: q.answers.all()
        
        answers_data = []
        for ans in q_answers:
            answers_data.append({
                'answer_text': ans.answer_text,
                'answer_created_at': ans.created_at.isoformat(),
                'answer_user': ans.user.username,
            })
        questions_data.append({
            'question_text': q.question_text,
            'question_created_at': q.created_at.isoformat(),
            'answers': answers_data
        })

    final_data = {
        'username': target_user.username,
        'questions': questions_data,
    }

    # Python'da sözlüğü JSON stringine dönüştür
    json_string = json.dumps(
        final_data,
        ensure_ascii=False,
        indent=2
    )

    # HttpResponse ile "indirilebilir" döndür
    response = HttpResponse(json_string, content_type='application/json; charset=utf-8')
    response['Content-Disposition'] = 'attachment; filename="entries.json"'
    return response

def download_entries_xlsx(request, username):
    """
    Kullanıcının oluşturduğu tüm soruları ve bu sorulara
    kullanıcının verdiği yanıtları Excel formatında indirir.

    Format:
    - A sütununda soru metni (her soru ilk boş satırda)
    - B sütununda yanıtlar (her yanıt bir alt satırda)
    - Sonraki soru, son yanıtın altındaki satırdan itibaren başlar.
    """

    # 1) Kullanıcıyı doğrula
    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse(
            {'error': 'Bu işlemi yapmaya yetkiniz yok.'},
            status=403
        )

    # 2) Kullanıcının oluşturduğu soruları çek
    user_questions = Question.objects.filter(user=target_user).order_by('created_at')

    # 3) Excel çalışma kitabı ve sayfa oluştur
    wb = Workbook()
    ws = wb.active
    ws.title = "Entries"

    # (İsteğe Bağlı) Başlık Satırı Ekleyebilirsiniz
    # ws["A1"] = "Soru"
    # ws["B1"] = "Yanıt"
    # ws["A1"].font = Font(bold=True)
    # ws["B1"].font = Font(bold=True)
    #
    # row_idx = 2  # Başlık satırı kullanıyorsak 2'den başlayacağız
    #
    # Aşağıda örnek, başlık satırı yokmuş gibi direkt 1. satırdan başlıyoruz:

    row_idx = 1

    # 4) Soruları gezip, satır-sütun mantığıyla verileri yerleştirelim
    for question in user_questions:
        # A sütununa soru metnini yaz
        ws.cell(row=row_idx, column=1, value=question.question_text)

        # Bu soruya kullanıcının verdiği tüm yanıtları çek
        # (Eğer tüm kullanıcılardan gelen yanıtları istiyorsanız: question.answers.all())
        user_answers = question.answers.filter(user=target_user).order_by('created_at')

        # Yanıtları B sütunundan itibaren ekleyelim
        # İlk yanıt, aynı satırda (B sütunu) olacak
        # Sonraki yanıtlar bir alt satıra inerek (B sütununda)
        answer_start_row = row_idx  # Soru metnini koyduğumuz satır

        for i, ans in enumerate(user_answers):
            # İlk yanıt aynı satırda, sonrakiler bir alt satırda
            # eğer her yanıtı ayrı satıra koymak istiyorsak:
            current_row = answer_start_row + i
            ws.cell(row=current_row, column=2, value=ans.answer_text)
        
        # Kaç tane yanıt varsa, bir sonraki soru
        # son yanıtın altındaki satırdan başlasın
        row_idx = answer_start_row + max(len(user_answers), 1)  # Yanıt yoksa da 1 satır ilerlet
        row_idx += 1  # 1 satır boşluk bırakmak isterseniz
        
    # 5) HttpResponse ile Excel'i "attachment" olarak döndür
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    # Dosya ismini dilediğiniz gibi oluşturabilirsiniz
    response['Content-Disposition'] = 'attachment; filename="entries.xlsx"'

    wb.save(response)
    return response

@login_required
def filter_answers(request, question_id):
    """
    Ajax filtre endpoint'i.
    Soruya ait yanıtları, my_answers / followed / username / keyword'e göre süzer
    ve 'core/_answers_list.html' partial'ını döndürür.
    """
    question = get_object_or_404(Question, id=question_id)
    
    # Parametreler
    my_answers = request.GET.get('my_answers')
    followed = request.GET.get('followed')
    username = request.GET.get('username', '').strip()
    keyword = request.GET.get('keyword', '').strip()

    # Tüm yanıtlar (bu soru altındaki)
    answers = question.answers.all().order_by('-created_at')  # güncelden eskiye

    # 1) Kendi yanıtlarım
    if my_answers == 'on':
        answers = answers.filter(user=request.user)

    # 2) Takip ettiklerim
    if followed == 'on':
        # UserProfile.following => diğer UserProfile'lar
        # Bu profillerin user'larına ait yanıtlar
        user_profile = request.user.userprofile
        followed_profiles = user_profile.following.all()
        # Yöntem 2: Liste halinde user'ları çıkarıyoruz
        followed_users = [p.user for p in followed_profiles]
        answers = answers.filter(user__in=followed_users)
    
    # 3) Kullanıcı adı (kısmi eşleşme)
    if username:
        # 'iexact' tam eşleşme yerine 'icontains' => kısmi
        answers = answers.filter(user__username__icontains=username)
    
    # 4) Kelime arama
    if keyword:
        answers = answers.filter(answer_text__icontains=keyword)

    # Kaydetme/Oylama bilgileri
    content_type_answer = ContentType.objects.get_for_model(Answer)
    answer_ids = answers.values_list('id', flat=True)

    # Kaydedilme sayıları
    saved_items = SavedItem.objects.filter(
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values('object_id').annotate(count=Count('id'))
    answer_save_dict = {item['object_id']: item['count'] for item in saved_items}

    # Kullanıcının kaydettiği yanıtlar
    saved_answer_ids = SavedItem.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values_list('object_id', flat=True)

    # Kullanıcının oy bilgisi
    user_votes = Vote.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values('object_id', 'value')
    user_vote_dict = {v['object_id']: v['value'] for v in user_votes}

    # up/down hesaplama
    for ans in answers:
        ans.user_vote_value = user_vote_dict.get(ans.id, 0)
        ans.upvotes = Vote.objects.filter(object_id=ans.id, value=1).count()
        ans.downvotes = Vote.objects.filter(object_id=ans.id, value=-1).count()

    # partial HTML döndür
    html_content = render_to_string(
        'core/_answers_list.html',
        {
            'answers': answers,
            'question': question,
            'answer_save_dict': answer_save_dict,
            'saved_answer_ids': saved_answer_ids,
            'search_keyword': keyword,
        },
        request=request
    )
    return HttpResponse(html_content)


def insert_toc(paragraph):
    # İçindekiler tablosu alanını ekler (Word'de elle güncellenir)
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_question_tree_to_docx(doc, question, target_user, level=1, visited=None):
    if visited is None:
        visited = set()
    if question.id in visited:
        return
    visited.add(question.id)

    doc.add_heading(question.question_text, level=level)

    user_answers = question.answers.filter(user=target_user).order_by('created_at')
    for answer in user_answers:
        date_str = answer.created_at.strftime("%Y-%m-%d %H:%M")
        p = doc.add_paragraph()
        run = p.add_run(date_str + "  ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(140, 140, 140)
        run.italic = True
        p.add_run(answer.answer_text)
        doc.add_paragraph("")

    subquestions = question.subquestions.all().order_by('created_at')
    for sub in subquestions:
        add_question_tree_to_docx(doc, sub, target_user, level=min(level+1, 9), visited=visited)

@login_required
def download_entries_docx(request, username):
    

    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse({'error': 'Bu işlemi yapmaya yetkiniz yok.'}, status=403)
    
    # Kök (üstte parent'ı olmayan) sorular
    user_questions = Question.objects.filter(user=target_user, parent_questions__isnull=True).order_by('created_at').distinct()

    document = Document()

    # Kullanıcı başlığı
    document.add_heading(f"{target_user.username} Entries", 0)

    # TOC ekle
    toc_paragraph = document.add_paragraph()
    insert_toc(toc_paragraph)

    # Talimat paragrafı
    instruction_text = (
        "Belgeyi açtıktan sonra içindekiler bölümünü görmek için, Word içerisinde "
        "alanı (veya tüm belgeyi) güncellemeniz gerekir (sağ tıklayıp 'Update Field' veya Ctrl+A ardından F9'a basabilirsiniz)."
    )
    document.add_paragraph(instruction_text)

    # Sayfa sonu
    document.add_page_break()

    # Hiyerarşik şekilde tüm başlıkları ve alt başlıkları recursive olarak ekle
    for q in user_questions:
        add_question_tree_to_docx(document, q, target_user, level=1)

    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(
        f.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{target_user.username}_entries.docx"'
    return response




def game_of_life(request):
    return render(request, "core/game_of_life.html")


def iat_start(request):
    return render(request, 'core/iat_start.html')

@login_required
def iat_test(request):
    # test_type: "gender" veya "ethnicity"
    test_type = request.GET.get('test_type')
    if test_type not in ["gender", "ethnicity"]:
        return redirect('iat_start')
    # Bu sayfada gerçek testi başlatacağız
    return render(request, 'core/iat_test.html', {'test_type': test_type})


@login_required
def iat_result_page(request):
    data = request.session.get('iat_result')
    if not data:
        return redirect("iat_start")

    if data['test_type'] == 'gender':
        reading = [
            {"title": "IAT nedir? (Bilinçdışı cinsiyet çağrışımı)", "url": "https://implicit.harvard.edu/implicit/takeatest.html"},
            {"title": "Toplumsal cinsiyet rolleri ve önyargı (makale)", "url": "https://psycnet.apa.org/record/2002-11021-009"},
            {"title": "Kadın-Erkek IAT analizleri", "url": "https://www.apa.org/research/action/sex-differences"},
        ]
    else:
        reading = [
            {"title": "Etnik önyargı ve çağrışımlar (IAT)", "url": "https://implicit.harvard.edu/implicit/takeatest.html"},
            {"title": "Türkiye'de Etnisite ve Önyargı", "url": "https://dergipark.org.tr/tr/download/article-file/224166"},
            {"title": "Kürt ve Türk kimliklerinde sosyal psikoloji", "url": "https://www.sciencedirect.com/science/article/pii/S0191886911003717"},
        ]

    ranking = data.get("ranking", [])
    return render(request, "core/iat_result.html", {
        "result": data,
        "reading": reading,
        "ranking": ranking,
    })



 # Sonradan model ekleyeceğiz


@login_required
@csrf_exempt
def iat_result(request):
    if request.method == "POST":
        data = json.loads(request.body)
        responses = data.get('responses', [])
        test_type = data.get('test_type', 'gender')

        # Aşamaları ayır: (0,2) = bir kategoriye karşı diğeri; (1,3) = kombinasyon eşleşmeleri
        stage_times = {i: [] for i in range(4)}
        stage_errors = {i: 0 for i in range(4)}
        for resp in responses:
            stage = resp['stage']
            rt = resp['rt']
            correct = resp['correct']
            if not correct:
                rt += 600   # Yanlış cevaplara 600 ms ekle
                stage_errors[stage] += 1
            stage_times[stage].append(rt)

        # D-score (Greenwald, 2003)
        # Anahtar bloklar: 2 ve 3 (kombinasyonlar)
        def safe_mean(x):
            return sum(x) / len(x) if x else 0
        def pooled_std(a, b):
            from math import sqrt
            aa = a if a else [1]
            bb = b if b else [1]
            mean_a = safe_mean(aa)
            mean_b = safe_mean(bb)
            var_a = sum((v - mean_a) ** 2 for v in aa) / (len(aa) - 1 if len(aa) > 1 else 1)
            var_b = sum((v - mean_b) ** 2 for v in bb) / (len(bb) - 1 if len(bb) > 1 else 1)
            return ((var_a + var_b) / 2) ** 0.5

        rt2 = [rt for rt in stage_times[2] if 200 < rt < 3000]
        rt3 = [rt for rt in stage_times[3] if 200 < rt < 3000]
        mean2 = safe_mean(rt2)
        mean3 = safe_mean(rt3)
        std_pooled = pooled_std(rt2, rt3) or 1
        dscore = (mean3 - mean2) / std_pooled

        # Yorum ve etiket
        if abs(dscore) < 0.15:
            bias_label = "Belirgin bir önyargı/çağrışım tespit edilmedi."
            bias_side = "neutral"
        elif dscore > 0.15:
            if test_type == "gender":
                bias_label = "Kadınları iyiyle, erkekleri kötüyle daha hızlı eşlediniz (Kadın lehine çağrışım)."
            else:
                bias_label = "Türkleri iyiyle, Kürtleri kötüyle daha hızlı eşlediniz (Türk lehine çağrışım)."
            bias_side = "left"
        else:
            if test_type == "gender":
                bias_label = "Erkekleri iyiyle, kadınları kötüyle daha hızlı eşlediniz (Erkek lehine çağrışım)."
            else:
                bias_label = "Kürtleri iyiyle, Türkleri kötüyle daha hızlı eşlediniz (Kürt lehine çağrışım)."
            bias_side = "right"

        # Sonucu kaydet
        result_obj = IATResult.objects.create(
            user=request.user,
            test_type=test_type,
            dscore=round(dscore, 3),
            bias_side=bias_side,
            errors=sum(stage_errors.values()),
            avg_rt=round((mean2 + mean3) / 2 if (mean2 + mean3) > 0 else 0)
        )

        # Sıralama listesi (aynı test tipine göre, yüksekten düşüğe)
        all_results = IATResult.objects.filter(test_type=test_type).order_by('-dscore')
        ranking = []
        user_rank = None
        for idx, res in enumerate(all_results, 1):
            is_me = (res.user == request.user)
            if is_me:
                user_rank = idx
            ranking.append({
                "rank": idx,
                "username": res.user.username if is_me else f"Kullanıcı #{idx}",
                "dscore": res.dscore,
                "is_me": is_me,
            })

        # Session ile sonucu ve sıralamayı gönder
        request.session['iat_result'] = {
            "test_type": test_type,
            "dscore": round(dscore, 3),
            "label": bias_label,
            "errors": sum(stage_errors.values()),
            "avg_rt": round((mean2 + mean3) / 2 if (mean2 + mean3) > 0 else 0),
            "bias_side": bias_side,
            "user_rank": user_rank,
            "ranking": ranking,
        }
        return JsonResponse({"result_url": reverse("iat_result_page")})

    return HttpResponse(status=405)


@csrf_exempt
@login_required
def kenarda_save(request):
    if request.method == "POST":
        data = json.loads(request.body)
        question_id = data.get("question_id")
        content = data.get("content")
        question = None
        if question_id:
            try:
                question = Question.objects.get(id=question_id)
            except Question.DoesNotExist:
                question = None
        # Mevcut taslak var mı?
        existing = Kenarda.objects.filter(user=request.user, question=question, is_sent=False).first()
        if existing:
            existing.content = content
            existing.updated_at = timezone.now()
            existing.save()
        else:
            Kenarda.objects.create(
                user=request.user,
                question=question,
                content=content,
                is_sent=False
            )
        return JsonResponse({"status": "ok"})
    return JsonResponse({"status": "fail", "error": "Yetkisiz veya geçersiz istek"})


@login_required
def kenarda_list(request):
    taslaklar = Kenarda.objects.filter(user=request.user, is_sent=False).order_by('-updated_at')
    return render(request, 'core/kenarda_list.html', {'taslaklar': taslaklar})




@require_POST
@login_required
def kenarda_sil(request, pk):
    try:
        taslak = Kenarda.objects.get(pk=pk, user=request.user)
        taslak.delete()
        return JsonResponse({"status":"ok"})
    except Kenarda.DoesNotExist:
        return JsonResponse({"status":"fail"})

@require_POST
@login_required
def kenarda_gonder(request, pk):
    try:
        taslak = Kenarda.objects.get(pk=pk, user=request.user)
        # Otomatik yanıt oluştur
        Answer.objects.create(
            question=taslak.question,
            answer_text=taslak.content,
            user=request.user
        )
        taslak.delete()
        return JsonResponse({"status":"ok"})
    except Exception as e:
        return JsonResponse({"status":"fail", "error": str(e)})


#ÇIKIŞ TESTLERİ
@login_required
def cikis_testleri_list(request):
    testler = CikisTesti.objects.filter(owner=request.user).order_by('-created_at')
    return render(request, 'core/cikis_testleri_list.html', {'testler': testler})

@login_required
def cikis_testi_create(request):
    if request.method == 'POST':
        form = CikisTestiForm(request.POST)
        if form.is_valid():
            test = form.save(commit=False)
            test.owner = request.user
            test.save()
            return redirect('cikis_testi_detail', test_id=test.id)
    else:
        form = CikisTestiForm()
    return render(request, 'core/cikis_testi_create.html', {'form': form})

@login_required
def cikis_testi_detail(request, test_id):
    test = get_object_or_404(CikisTesti, id=test_id, owner=request.user)
    sorular = test.sorular.all().order_by('order', 'id')
    return render(request, 'core/cikis_testi_detail.html', {'test': test, 'sorular': sorular})

@login_required
def cikis_soru_ekle(request, test_id):
    test = get_object_or_404(CikisTesti, id=test_id, owner=request.user)
    if request.method == 'POST':
        form = CikisTestiSoruForm(request.POST)
        if form.is_valid():
            soru = form.save(commit=False)
            soru.test = test
            soru.save()
            return redirect('cikis_testi_detail', test_id=test.id)
    else:
        form = CikisTestiSoruForm()
    return render(request, 'core/cikis_soru_ekle.html', {'form': form, 'test': test})

@login_required
def cikis_sik_ekle(request, soru_id):
    soru = get_object_or_404(CikisTestiSoru, id=soru_id, test__owner=request.user)
    if request.method == 'POST':
        form = CikisTestiSikForm(request.POST)
        if form.is_valid():
            sik = form.save(commit=False)
            sik.soru = soru
            sik.save()
            return redirect('cikis_testi_detail', test_id=soru.test.id)
    else:
        form = CikisTestiSikForm()
    return render(request, 'core/cikis_sik_ekle.html', {'form': form, 'soru': soru})

@login_required
def cikis_dogru_sik_sec(request, soru_id):
    soru = get_object_or_404(CikisTestiSoru, id=soru_id, test__owner=request.user)
    if request.method == 'POST':
        sik_id = request.POST.get('correct_option')
        try:
            sik = soru.siklar.get(id=sik_id)
            soru.correct_option = sik
            soru.save()
        except CikisTestiSik.DoesNotExist:
            pass
        return redirect('cikis_testi_detail', test_id=soru.test.id)
    siklar = soru.siklar.all()
    return render(request, 'core/cikis_dogru_sik_sec.html', {'soru': soru, 'siklar': siklar})



@login_required
def cikis_testi_coz(request, test_id):
    test = get_object_or_404(CikisTesti, id=test_id)
    sorular = test.sorular.all()
    if request.method == "POST":
        cevaplar = {}
        dogru = 0
        for soru in sorular:
            secilen = request.POST.get(f"soru_{soru.id}")
            cevaplar[soru.id] = secilen
            if secilen and soru.correct_option and int(secilen) == soru.correct_option.id:
                dogru += 1
        # Sonucu kaydet
        result = CikisTestiResult.objects.create(
            test=test,
            user=request.user,
            dogru_sayisi=dogru,
            toplam_soru=sorular.count(),
            cikis_dogrusu_uydu=(test.cikis_dogrusu is not None and dogru >= test.cikis_dogrusu)
        )
        return render(request, "core/cikis_testi_sonuc.html", {
            "test": test,
            "dogru": dogru,
            "toplam": sorular.count(),
            "cikis_dogrusu": test.cikis_dogrusu,
            "uydu": (test.cikis_dogrusu is not None and dogru >= test.cikis_dogrusu),
            "sonuc": result,
        })
    return render(request, "core/cikis_testi_coz.html", {"test": test, "sorular": sorular})

@login_required
def cikis_testi_sonuc_list(request, test_id):
    test = get_object_or_404(CikisTesti, id=test_id)
    # Doğru sayısına göre azalan, eşitse en güncel önde
    sonuclar = test.sonuclar.select_related('user').all().order_by('-dogru_sayisi', '-completed_at')
    return render(request, "core/cikis_testi_sonuc_list.html", {
        "test": test,
        "sonuclar": sonuclar,
        "cikis_dogrusu": test.cikis_dogrusu
    })


@login_required
def cikis_dogrusu_ayarla(request, test_id):
    test = get_object_or_404(CikisTesti, id=test_id, owner=request.user)
    if request.method == "POST":
        try:
            yeni_dogruluk = int(request.POST.get("cikis_dogrusu"))
            if 0 <= yeni_dogruluk <= test.sorular.count():
                test.cikis_dogrusu = yeni_dogruluk
                test.save()
        except (ValueError, TypeError):
            pass
        return redirect("cikis_testi_detail", test_id=test.id)
    # Detaydan çağırıyorsan burası çalışmamalı!
    return redirect("cikis_testi_detail", test_id=test.id)

@login_required
def cikis_soru_edit(request, soru_id):
    soru = get_object_or_404(CikisTestiSoru, id=soru_id, test__owner=request.user)
    if request.method == "POST":
        form = CikisTestiSoruForm(request.POST, instance=soru)
        if form.is_valid():
            form.save()
            return redirect('cikis_testi_detail', test_id=soru.test.id)
    else:
        form = CikisTestiSoruForm(instance=soru)
    return render(request, 'core/cikis_soru_edit.html', {'form': form, 'soru': soru})

@login_required
def cikis_soru_sil(request, soru_id):
    soru = get_object_or_404(CikisTestiSoru, id=soru_id, test__owner=request.user)
    test_id = soru.test.id
    if request.method == "POST":
        soru.delete()
        return redirect('cikis_testi_detail', test_id=test_id)
    return render(request, 'core/cikis_soru_sil.html', {'soru': soru})


def cikis_test_list(request):
    tests = CikisTesti.objects.all().order_by('-created_at')
    return render(request, 'core/cikis_test_list.html', {'tests': tests})

@login_required
def cikis_test_coz(request, test_id):
    test = get_object_or_404(CikisTesti, id=test_id)
    sorular = test.sorular.all()
    # Kullanıcı daha önce çözdüyse sonucu göster, yoksa cevaplamasını sağla
    if request.method == "POST":
        # Yanıtları kaydet ve sonucu hesapla
        # ...
        return redirect('cikis_test_sonuc', test_id=test.id)
    return render(request, 'core/cikis_test_coz.html', {'test': test, 'sorular': sorular})


@login_required
def cikis_sonuc_sil(request, sonuc_id):
    sonuc = get_object_or_404(CikisTestiResult, id=sonuc_id)
    if sonuc.user != request.user:
        return HttpResponseForbidden("Bunu silemezsiniz.")
    if request.method == "POST":
        sonuc.delete()
        # AJAX için uygun döndür
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"status": "ok"})
        # Normal redirect (eski usül link tıklandıysa)
        return redirect("cikis_testi_sonuc_list", test_id=sonuc.test.id)
    return HttpResponseForbidden("Geçersiz istek")