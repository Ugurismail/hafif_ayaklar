"""
Miscellaneous views
- get_today_questions_queryset
- download_entries_json
- download_entries_xlsx
- download_entries_docx
- filter_answers
- insert_toc
- add_question_tree_to_docx
"""

import json
import os
import random
import re
from collections import Counter
from datetime import timedelta
from io import BytesIO
from itertools import chain
from uuid import uuid4

from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.contrib.auth.models import User
from django.contrib.contenttypes.models import ContentType
from django.core.exceptions import PermissionDenied, SuspiciousOperation
from django.core.files.storage import default_storage
from django.db import DatabaseError
from django.db.models import Q, Count, Max, F
from django.db.models.functions import Coalesce, TruncWeek, TruncMonth
from django.http import JsonResponse, HttpResponse, Http404
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.utils.timezone import now
from django.views.decorators.http import require_POST
from PIL import Image, UnidentifiedImageError

from ..middleware import LastSeenMiddleware

from ..models import (
    Question, Answer, Vote, SavedItem, StartingQuestion,
    Reference, UserProfile, QuestionRelationship, LibraryFile, DailyVisitor
)
from ..utils import paginate_queryset, build_reference_usage_counts
from ..forms import LibraryFileForm
from ..services import VoteSaveService
from ..querysets import get_today_questions_queryset, get_active_left_frame_pin_q
from .answer_views import get_all_descendant_question_ids
from ..answer_git import attach_answer_revision_metadata

WORD_PATTERN = re.compile(r'\b\w+\b', re.UNICODE)
ALLOWED_EDITOR_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}
ALLOWED_EDITOR_IMAGE_MIME_TYPES = {'image/jpeg', 'image/png', 'image/gif', 'image/webp'}
MAX_EDITOR_IMAGE_SIZE = 5 * 1024 * 1024








@login_required




@login_required


@login_required
@require_POST


@login_required


@login_required


@login_required
@require_POST
































@require_POST




def collect_user_bibliography(target_user, specific_answers=None):
    """
    Kullanıcının entry'lerinden kaynakları toplar ve bibliyografya listesi döndürür.
    Args:
        target_user: User object
        specific_answers: Optional list/queryset of specific Answer objects to collect from
    Returns: List of dicts containing reference information with numbering
    """
    from core.models import Answer, Reference

    # Kullanıcının yanıtlarını al (belirtilmişse sadece belirli yanıtlar)
    if specific_answers is not None:
        user_answers = specific_answers
    else:
        user_answers = Answer.objects.filter(user=target_user)

    # Tüm entry metinlerini birleştir
    all_text = ' '.join([answer.answer_text for answer in user_answers if answer.answer_text])

    # extract_bibliography mantığını kullanarak kaynakları topla
    if not all_text:
        return []

    reference_ids = set()  # Store unique reference IDs
    reference_pages = {}  # Store page numbers for each reference

    # Support both long and short forms:
    # (kaynak:12, sayfa:10-12) / (k:12, s:10-12)
    # Allow optional comma + flexible spacing.
    pattern = re.compile(
        r'\((?:kaynak|k)\s*:\s*(\d+)(?:(?:\s*,?\s*(?:sayfa|s)\s*:\s*([^)]+)))?\)',
        re.IGNORECASE
    )
    matches = pattern.finditer(all_text)

    for match in matches:
        ref_id_str = match.group(1)
        sayfa = match.group(2)  # Optional: None or string (12-14, 123a etc.)
        ref_id = int(ref_id_str)

        reference_ids.add(ref_id)

        if ref_id not in reference_pages:
            reference_pages[ref_id] = []

        # Collect page numbers if they exist
        if sayfa:
            page_str = sayfa.strip()
            if page_str not in reference_pages[ref_id]:
                reference_pages[ref_id].append(page_str)

    # Build the bibliography list - sorted by reference ID
    bibliography = []
    for ref_id in sorted(reference_ids):
        try:
            ref_obj = Reference.objects.get(id=ref_id)
            pages = reference_pages.get(ref_id, [])

            # Çoklu yazarları düzgün formatla
            surnames = [s.strip() for s in ref_obj.author_surname.split(';') if s.strip()]
            names = [n.strip() for n in ref_obj.author_name.split(';') if n.strip()]

            authors = []
            for i in range(max(len(surnames), len(names))):
                surname = surnames[i] if i < len(surnames) else ''
                name = names[i] if i < len(names) else ''
                if surname or name:
                    authors.append(f"{surname}, {name}".strip(', '))

            formatted_authors = '; '.join(authors)

            bibliography.append({
                'number': ref_id,  # Orijinal kaynak ID'sini kullan
                'reference': ref_obj,
                'formatted_authors': formatted_authors,
                'pages': pages
            })
        except Reference.DoesNotExist:
            bibliography.append({
                'number': ref_id,  # Orijinal kaynak ID'sini kullan
                'reference': None,
                'ref_id': ref_id,
                'pages': []
            })

    return bibliography


def get_filtered_user_answers(request, target_user):
    """
    Helper function to get filtered user answers based on request parameters.
    Handles entry_ids, root_question_id, and order parameters.
    Returns a queryset or list of Answer objects.
    """
    entry_ids = None
    order = 'oldest'  # default
    root_question_id = None

    if request.method == 'POST':
        entry_ids_str = request.POST.get('entry_ids', '')
        if entry_ids_str:
            entry_ids = [int(id.strip()) for id in entry_ids_str.split(',') if id.strip()]
        order = request.POST.get('order', 'oldest')
        root_question_id = request.POST.get('root_question_id', '')

    # Start with base queryset
    user_answers = Answer.objects.filter(user=target_user)

    # Filter by root question if provided
    if root_question_id:
        try:
            root_id = int(root_question_id)
            question_ids = get_all_descendant_question_ids(root_id, target_user)
            user_answers = user_answers.filter(question_id__in=question_ids)
        except (ValueError, TypeError):
            pass  # Invalid root_question_id, ignore

    # Filter by specific entry IDs if provided
    if entry_ids:
        user_answers = user_answers.filter(id__in=entry_ids)

    # Apply ordering
    if order == 'newest':
        user_answers = user_answers.order_by('-created_at')
    elif order == 'oldest':
        user_answers = user_answers.order_by('created_at')
    elif order == 'custom' and entry_ids:
        # For custom order, preserve the order from entry_ids list
        user_answers = list(user_answers)
        answers_dict = {ans.id: ans for ans in user_answers}
        user_answers = [answers_dict[id] for id in entry_ids if id in answers_dict]
    else:
        # Default to oldest
        user_answers = user_answers.order_by('created_at')

    return user_answers


def is_custom_order_request(request):
    """
    Custom order only makes sense when specific entry_ids are provided.
    """
    if request.method != 'POST':
        return False
    order = request.POST.get('order', 'oldest')
    entry_ids_str = (request.POST.get('entry_ids', '') or '').strip()
    return order == 'custom' and bool(entry_ids_str)


@login_required
def download_entries_json(request, username):
    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse(
            {'error': 'Bu işlemi yapmaya yetkiniz yok.'},
            status=403
        )

    # Get filtered answers using helper function
    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    # Group answers by question
    questions_dict = {}
    for ans in user_answers:
        q = ans.question
        if q.id not in questions_dict:
            questions_dict[q.id] = {
                'question': q,
                'answers': []
            }
        questions_dict[q.id]['answers'].append(ans)

    questions_data = []
    for _, q_data in questions_dict.items():
        q = q_data['question']
        q_answers = q_data['answers']

        answers_data = []
        for ans in q_answers:
            answers_data.append({
                'answer_text': ans.answer_text,
                'answer_created_at': ans.created_at.isoformat(),
                'answer_user': ans.user.username,
            })
        questions_data.append({
            'question_text': q.question_text,
            'question_created_at': q.created_at.isoformat(),
            'answers': answers_data
        })

    # Preserve the exact order (manual/selection order) when requested.
    entries_data = []
    if custom_order:
        for ans in user_answers:
            entries_data.append({
                'question_text': ans.question.question_text,
                'question_slug': ans.question.slug,
                'answer_id': ans.id,
                'answer_text': ans.answer_text,
                'answer_created_at': ans.created_at.isoformat(),
            })

    # Collect bibliography from selected answers
    bibliography = collect_user_bibliography(target_user, user_answers)
    references_data = []
    for bib_item in bibliography:
        if bib_item.get('reference'):
            ref = bib_item['reference']
            ref_dict = {
                'number': bib_item['number'],
                'authors': bib_item['formatted_authors'],
                'year': ref.year,
                'title': ref.metin_ismi or '',
                'rest': ref.rest,
                'pages_used': bib_item['pages']
            }
            if ref.abbreviation:
                ref_dict['abbreviation'] = ref.abbreviation
            references_data.append(ref_dict)
        else:
            # Reference not found
            references_data.append({
                'number': bib_item['number'],
                'ref_id': bib_item.get('ref_id'),
                'error': 'Reference not found'
            })

    final_data = {
        'username': target_user.username,
        'questions': questions_data,
        'entries': entries_data,
        'references': references_data
    }

    json_string = json.dumps(
        final_data,
        ensure_ascii=False,
        indent=2
    )

    response = HttpResponse(json_string, content_type='application/json; charset=utf-8')
    response['Content-Disposition'] = 'attachment; filename="entries.json"'
    return response


@login_required
def download_entries_xlsx(request, username):
    from openpyxl import Workbook

    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse(
            {'error': 'Bu işlemi yapmaya yetkiniz yok.'},
            status=403
        )

    # Get filtered answers using helper function
    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    wb = Workbook()
    ws = wb.active
    ws.title = "Entries"

    if custom_order:
        ws.cell(row=1, column=1, value="Soru")
        ws.cell(row=1, column=2, value="Tarih")
        ws.cell(row=1, column=3, value="Entry")
        for i, ans in enumerate(user_answers, start=2):
            ws.cell(row=i, column=1, value=ans.question.question_text)
            ws.cell(row=i, column=2, value=ans.created_at.strftime("%Y-%m-%d %H:%M"))
            ws.cell(row=i, column=3, value=ans.answer_text)
    else:
        # Group answers by question
        questions_dict = {}
        for ans in user_answers:
            q = ans.question
            if q.id not in questions_dict:
                questions_dict[q.id] = {
                    'question': q,
                    'answers': []
                }
            questions_dict[q.id]['answers'].append(ans)

        row_idx = 1
        for _, q_data in questions_dict.items():
            question = q_data['question']
            q_answers = q_data['answers']

            ws.cell(row=row_idx, column=1, value=question.question_text)

            answer_start_row = row_idx
            for j, ans in enumerate(q_answers):
                current_row = answer_start_row + j
                ws.cell(row=current_row, column=2, value=ans.answer_text)

            row_idx = answer_start_row + max(len(q_answers), 1)
            row_idx += 1

    # Add bibliography sheet
    bibliography = collect_user_bibliography(target_user, user_answers)
    if bibliography:
        ws_refs = wb.create_sheet(title="Kaynaklar")

        # Add headers
        ws_refs.cell(row=1, column=1, value="No")
        ws_refs.cell(row=1, column=2, value="Yazarlar")
        ws_refs.cell(row=1, column=3, value="Yıl")
        ws_refs.cell(row=1, column=4, value="Metin İsmi")
        ws_refs.cell(row=1, column=5, value="Künye")
        ws_refs.cell(row=1, column=6, value="Kullanılan Sayfalar")

        # Add bibliography items
        for idx, bib_item in enumerate(bibliography, start=2):
            if bib_item.get('reference'):
                ref = bib_item['reference']
                ws_refs.cell(row=idx, column=1, value=bib_item['number'])
                ws_refs.cell(row=idx, column=2, value=bib_item['formatted_authors'])
                ws_refs.cell(row=idx, column=3, value=ref.year)
                ws_refs.cell(row=idx, column=4, value=ref.metin_ismi or '')
                ws_refs.cell(row=idx, column=5, value=ref.rest)
                ws_refs.cell(row=idx, column=6, value=', '.join(bib_item['pages']) if bib_item['pages'] else '')
            else:
                ws_refs.cell(row=idx, column=1, value=bib_item['number'])
                ws_refs.cell(row=idx, column=2, value='Kaynak bulunamadı')

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="entries.xlsx"'

    wb.save(response)
    return response


def insert_toc(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_answer_text_to_docx(document, answer_text):
    """
    Add answer text as real Word paragraphs.
    - Double newlines create new paragraphs.
    - Single newlines are preserved as line breaks inside a paragraph.
    """
    if not answer_text:
        return

    normalized = str(answer_text).replace('\r\n', '\n').replace('\r', '\n').strip()
    if not normalized:
        return

    blocks = re.split(r'\n{2,}', normalized)
    for block in blocks:
        block = (block or '').strip('\n')
        if not block.strip():
            continue

        paragraph = document.add_paragraph()
        lines = block.split('\n')
        for idx, line in enumerate(lines):
            run = paragraph.add_run(line)
            if idx < len(lines) - 1:
                run.add_break()


def add_question_tree_to_docx(doc, question, target_user, level=1, visited=None):
    from docx.shared import Pt, RGBColor

    if visited is None:
        visited = set()
    if question.id in visited:
        return
    visited.add(question.id)

    doc.add_heading(question.question_text, level=level)

    user_answers = question.answers.filter(user=target_user).order_by('created_at')
    for answer in user_answers:
        date_str = answer.created_at.strftime("%Y-%m-%d %H:%M")
        p = doc.add_paragraph()
        run = p.add_run(date_str + "  ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(140, 140, 140)
        run.italic = True
        add_answer_text_to_docx(doc, answer.answer_text)
        doc.add_paragraph("")

    # Kullanıcının bu sorunun alt sorularını al (kullanıcı-bazlı)
    subquestions_rels = QuestionRelationship.objects.filter(
        parent=question,
        user=target_user
    ).select_related('child').order_by('created_at')
    for rel in subquestions_rels:
        add_question_tree_to_docx(doc, rel.child, target_user, level=min(level+1, 9), visited=visited)


@login_required
def download_entries_docx(request, username):
    from docx import Document
    from docx.shared import Pt, RGBColor

    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse({'error': 'Bu işlemi yapmaya yetkiniz yok.'}, status=403)

    # Get filtered answers using helper function
    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    document = Document()

    document.add_heading(f"{target_user.username} Entries", 0)

    if custom_order:
        last_question_id = None
        for ans in user_answers:
            if ans.question_id != last_question_id:
                document.add_heading(ans.question.question_text, level=1)
                last_question_id = ans.question_id

            date_str = ans.created_at.strftime("%Y-%m-%d %H:%M")
            p = document.add_paragraph()
            run = p.add_run(date_str + "  ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(140, 140, 140)
            run.italic = True
            add_answer_text_to_docx(document, ans.answer_text)
            document.add_paragraph("")
    else:
        # Group answers by question
        questions_dict = {}
        for ans in user_answers:
            q = ans.question
            if q.id not in questions_dict:
                questions_dict[q.id] = {
                    'question': q,
                    'answers': []
                }
            questions_dict[q.id]['answers'].append(ans)

        toc_paragraph = document.add_paragraph()
        insert_toc(toc_paragraph)

        instruction_text = (
            "Belgeyi açtıktan sonra içindekiler bölümünü görmek için, Word içerisinde "
            "alanı (veya tüm belgeyi) güncellemeniz gerekir (sağ tıklayıp 'Update Field' veya Ctrl+A ardından F9'a basabilirsiniz)."
        )
        document.add_paragraph(instruction_text)

        document.add_page_break()

        for _, q_data in questions_dict.items():
            question = q_data['question']
            q_answers = q_data['answers']

            document.add_heading(question.question_text, level=1)

            for answer in q_answers:
                date_str = answer.created_at.strftime("%Y-%m-%d %H:%M")
                p = document.add_paragraph()
                run = p.add_run(date_str + "  ")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(140, 140, 140)
                run.italic = True
                add_answer_text_to_docx(document, answer.answer_text)
                document.add_paragraph("")

    # Add bibliography section
    bibliography = collect_user_bibliography(target_user, user_answers)
    if bibliography:
        document.add_page_break()
        document.add_heading('Kaynakça', level=1)

        for bib_item in bibliography:
            if bib_item.get('reference'):
                ref = bib_item['reference']
                # Format: [1] Dumézil, Georges (1950), Metin İsmi, Künye
                ref_text = f"[{bib_item['number']}] {bib_item['formatted_authors']} ({ref.year})"
                if ref.metin_ismi:
                    ref_text += f", {ref.metin_ismi}"
                ref_text += f", {ref.rest}"

                # Add pages used if available
                if bib_item['pages']:
                    ref_text += f" (Kullanılan sayfalar: {', '.join(bib_item['pages'])})"

                p = document.add_paragraph(ref_text)
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            else:
                # Reference not found
                ref_text = f"[{bib_item['number']}] Kaynak bulunamadı (ID: {bib_item.get('ref_id')})"
                p = document.add_paragraph(ref_text)
                p.paragraph_format.left_indent = Pt(18)

    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(
        f.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{target_user.username}_entries.docx"'
    return response


@login_required
def download_entries_pdf(request, username):
    """
    PDF export for user entries using ReportLab
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse({'error': 'Bu işlemi yapmaya yetkiniz yok.'}, status=403)

    # Get filtered answers using helper function
    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    # Create PDF buffer
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)

    # Container for PDF elements
    elements = []

    # Register Turkish-compatible font (DejaVu Sans)
    try:
        import os
        from django.conf import settings
        import logging
        logger = logging.getLogger(__name__)

        # Use DejaVu Sans font from static/fonts directory
        font_dir = os.path.join(settings.BASE_DIR, 'static', 'fonts')
        dejavu_regular = os.path.join(font_dir, 'DejaVuSans.ttf')
        dejavu_bold = os.path.join(font_dir, 'DejaVuSans-Bold.ttf')

        logger.info(f"Font paths - Regular: {dejavu_regular}, Bold: {dejavu_bold}")
        logger.info(f"Regular exists: {os.path.exists(dejavu_regular)}, Bold exists: {os.path.exists(dejavu_bold)}")

        if os.path.exists(dejavu_regular) and os.path.exists(dejavu_bold):
            pdfmetrics.registerFont(TTFont('TurkishFont', dejavu_regular))
            pdfmetrics.registerFont(TTFont('TurkishFont-Bold', dejavu_bold))
            font_name = 'TurkishFont'
            font_name_bold = 'TurkishFont-Bold'
            logger.info("DejaVu Sans fonts successfully registered!")
        else:
            # Fallback to Helvetica if fonts not found
            logger.warning("DejaVu fonts not found! Falling back to Helvetica (no Turkish support)")
            font_name = 'Helvetica'
            font_name_bold = 'Helvetica-Bold'
    except Exception as e:
        # If font registration fails, use built-in Helvetica
        logger.error(f"Font registration failed: {e}")
        font_name = 'Helvetica'
        font_name_bold = 'Helvetica-Bold'

    # Define styles
    styles = getSampleStyleSheet()

    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=24,
        textColor='#2c3e50',
        spaceAfter=30,
        alignment=TA_CENTER
    )

    # Heading styles for different levels
    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=16,
        textColor='#2c3e50',
        spaceAfter=12,
        spaceBefore=12
    )

    h2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=14,
        textColor='#34495e',
        spaceAfter=10,
        spaceBefore=10,
        leftIndent=20
    )

    h3_style = ParagraphStyle(
        'CustomH3',
        parent=styles['Heading3'],
        fontName=font_name,
        fontSize=12,
        textColor='#7f8c8d',
        spaceAfter=8,
        spaceBefore=8,
        leftIndent=40
    )

    # Answer text style
    answer_style = ParagraphStyle(
        'AnswerText',
        parent=styles['BodyText'],
        fontName=font_name,
        fontSize=10,
        textColor='#2c3e50',
        spaceAfter=12,
        leftIndent=20
    )

    # Date style
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=8,
        textColor='#95a5a6',
        spaceAfter=6,
        leftIndent=20
    )

    # Add title
    title = Paragraph(f"{target_user.username} - Entry'ler", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.3*inch))

    # Helper function to escape HTML entities and handle special chars
    def clean_text(text):
        if not text:
            return ""
        # Replace HTML-like characters
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        # Handle line breaks
        text = text.replace('\n', '<br/>')
        return text

    if custom_order:
        last_question_id = None
        for answer in user_answers:
            if answer.question_id != last_question_id:
                question_para = Paragraph(clean_text(answer.question.question_text), h1_style)
                elements.append(question_para)
                elements.append(Spacer(1, 0.1*inch))
                last_question_id = answer.question_id

            date_str = answer.created_at.strftime("%Y-%m-%d %H:%M")
            date_para = Paragraph(f"<i>{date_str}</i>", date_style)
            elements.append(date_para)
            answer_para = Paragraph(clean_text(answer.answer_text), answer_style)
            elements.append(answer_para)
            elements.append(Spacer(1, 0.15*inch))
    else:
        # Group answers by question
        questions_dict = {}
        for ans in user_answers:
            q = ans.question
            if q.id not in questions_dict:
                questions_dict[q.id] = {
                    'question': q,
                    'answers': []
                }
            questions_dict[q.id]['answers'].append(ans)

        # Add Table of Contents
        toc_style = ParagraphStyle(
            'TOCHeading',
            parent=styles['Heading1'],
            fontName=font_name_bold,
            fontSize=18,
            textColor='#2c3e50',
            spaceAfter=12
        )
        toc_item_style = ParagraphStyle(
            'TOCItem',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            textColor='#34495e',
            spaceAfter=6,
            leftIndent=20
        )

        elements.append(Paragraph("İçindekiler", toc_style))
        elements.append(Spacer(1, 0.2*inch))

        for idx, (_, q_data) in enumerate(questions_dict.items(), 1):
            question = q_data['question']
            toc_text = f"{idx}. {clean_text(question.question_text[:100])}"
            if len(question.question_text) > 100:
                toc_text += "..."
            elements.append(Paragraph(toc_text, toc_item_style))

        elements.append(PageBreak())

        for _, q_data in questions_dict.items():
            question = q_data['question']
            q_answers = q_data['answers']

            question_para = Paragraph(clean_text(question.question_text), h1_style)
            elements.append(question_para)
            elements.append(Spacer(1, 0.1*inch))

            for answer in q_answers:
                date_str = answer.created_at.strftime("%Y-%m-%d %H:%M")
                date_para = Paragraph(f"<i>{date_str}</i>", date_style)
                elements.append(date_para)

                answer_para = Paragraph(clean_text(answer.answer_text), answer_style)
                elements.append(answer_para)
                elements.append(Spacer(1, 0.15*inch))

            elements.append(PageBreak())

    # Add bibliography section
    bibliography = collect_user_bibliography(target_user, user_answers)
    if bibliography:
        # Add page break before bibliography
        elements.append(PageBreak())

        # Bibliography heading style
        bib_heading_style = ParagraphStyle(
            'BibliographyHeading',
            parent=styles['Heading1'],
            fontName=font_name_bold,
            fontSize=18,
            textColor='#2c3e50',
            spaceAfter=20,
            spaceBefore=12
        )

        # Bibliography item style
        bib_item_style = ParagraphStyle(
            'BibliographyItem',
            parent=styles['BodyText'],
            fontName=font_name,
            fontSize=10,
            textColor='#2c3e50',
            spaceAfter=10,
            leftIndent=20,
            firstLineIndent=-20
        )

        # Add bibliography heading
        bib_heading = Paragraph("Kaynakça", bib_heading_style)
        elements.append(bib_heading)
        elements.append(Spacer(1, 0.2*inch))

        # Add bibliography items
        for bib_item in bibliography:
            if bib_item.get('reference'):
                ref = bib_item['reference']
                # Format: [1] Dumézil, Georges (1950), Metin İsmi, Künye
                ref_text = f"[{bib_item['number']}] {clean_text(bib_item['formatted_authors'])} ({ref.year})"
                if ref.metin_ismi:
                    ref_text += f", {clean_text(ref.metin_ismi)}"
                ref_text += f", {clean_text(ref.rest)}"

                # Add pages used if available
                if bib_item['pages']:
                    pages_str = ', '.join(bib_item['pages'])
                    ref_text += f" (Kullanılan sayfalar: {pages_str})"

                bib_para = Paragraph(ref_text, bib_item_style)
                elements.append(bib_para)
            else:
                # Reference not found
                ref_text = f"[{bib_item['number']}] Kaynak bulunamadı (ID: {bib_item.get('ref_id')})"
                bib_para = Paragraph(ref_text, bib_item_style)
                elements.append(bib_para)

    # Build PDF
    doc.build(elements)

    # Get PDF from buffer
    buffer.seek(0)
    pdf_content = buffer.read()

    # Create response
    response = HttpResponse(pdf_content, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{target_user.username}_entries.pdf"'
    return response


@login_required
def filter_answers(request, slug):
    """
    Ajax filtre endpoint'i.
    Soruya ait yanıtları, my_answers / followed / username / keyword'e göre süzer
    ve 'core/_answers_list.html' partial'ını döndürür.
    """
    question = get_object_or_404(Question, slug=slug)

    # Parametreler
    my_answers = request.GET.get('my_answers')
    followed = request.GET.get('followed')
    username = request.GET.get('username', '').strip()
    keyword = request.GET.get('keyword', '').strip()

    # Tüm yanıtlar (bu soru altındaki)
    answers = question.answers.all().order_by('created_at')

    # 1) Kendi yanıtlarım
    if my_answers == 'on':
        answers = answers.filter(user=request.user)

    # 2) Takip ettiklerim
    if followed == 'on':
        user_profile = request.user.userprofile
        followed_profiles = user_profile.following.all()
        followed_users = [p.user for p in followed_profiles]
        answers = answers.filter(user__in=followed_users)

    # 3) Kullanıcı adı (kısmi eşleşme)
    if username:
        answers = answers.filter(user__username__icontains=username)

    # 4) Kelime arama
    if keyword:
        answers = answers.filter(answer_text__icontains=keyword)

    # Kaydetme/Oylama bilgileri
    content_type_answer = ContentType.objects.get_for_model(Answer)
    answer_ids = answers.values_list('id', flat=True)

    # Kaydedilme sayıları
    saved_items = SavedItem.objects.filter(
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values('object_id').annotate(count=Count('id'))
    answer_save_dict = {item['object_id']: item['count'] for item in saved_items}

    # Kullanıcının kaydettiği yanıtlar
    saved_answer_ids = SavedItem.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values_list('object_id', flat=True)

    # Kullanıcının oy bilgisi
    user_votes = Vote.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values('object_id', 'value')
    user_vote_dict = {v['object_id']: v['value'] for v in user_votes}

    # up/down hesaplama
    answers = list(answers.select_related('user', 'question'))
    attach_answer_revision_metadata(answers, current_user=request.user)
    for ans in answers:
        ans.user_vote_value = user_vote_dict.get(ans.id, 0)
        ans.upvotes = Vote.objects.filter(object_id=ans.id, value=1).count()
        ans.downvotes = Vote.objects.filter(object_id=ans.id, value=-1).count()

    # partial HTML döndür
    html_content = render_to_string(
        'core/_answers_list.html',
        {
            'answers': answers,
            'question': question,
            'answer_save_dict': answer_save_dict,
            'saved_answer_ids': saved_answer_ids,
            'search_keyword': keyword,
        },
        request=request
    )
    return HttpResponse(html_content)
