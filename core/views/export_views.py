"""Export and filtered-answer related views."""

import json
import re
from io import BytesIO

from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.contenttypes.models import ContentType
from django.db.models import Count
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string

from ..answer_git import attach_answer_revision_metadata
from ..models import Answer, Question, QuestionRelationship, Reference, SavedItem, Vote
from .answer_views import get_all_descendant_question_ids


def collect_user_bibliography(target_user, specific_answers=None):
    """Collect bibliography for a user's selected answers."""
    from core.models import Answer, Reference

    if specific_answers is not None:
        user_answers = specific_answers
    else:
        user_answers = Answer.objects.filter(user=target_user)

    all_text = ' '.join([answer.answer_text for answer in user_answers if answer.answer_text])
    if not all_text:
        return []

    reference_ids = set()
    reference_pages = {}

    pattern = re.compile(
        r'\((?:kaynak|k)\s*:\s*(\d+)(?:(?:\s*,?\s*(?:sayfa|s)\s*:\s*([^)]+)))?\)',
        re.IGNORECASE
    )
    matches = pattern.finditer(all_text)

    for match in matches:
        ref_id_str = match.group(1)
        sayfa = match.group(2)
        try:
            ref_id = int(ref_id_str)
            reference_ids.add(ref_id)
            if sayfa:
                if ref_id not in reference_pages:
                    reference_pages[ref_id] = set()
                reference_pages[ref_id].add(sayfa.strip())
        except ValueError:
            continue

    bibliography = []
    for ref_id in sorted(reference_ids):
        try:
            ref_obj = Reference.objects.get(id=ref_id)
            pages = reference_pages.get(ref_id, [])

            surnames = [s.strip() for s in (ref_obj.author_surname or '').split(';') if s.strip()]
            names = [n.strip() for n in (ref_obj.author_name or '').split(';') if n.strip()]

            authors = []
            for i in range(max(len(surnames), len(names))):
                surname = surnames[i] if i < len(surnames) else ''
                name = names[i] if i < len(names) else ''
                if surname or name:
                    authors.append(f"{surname}, {name}".strip(', '))

            formatted_authors = '; '.join(authors)

            bibliography.append({
                'number': ref_id,
                'reference': ref_obj,
                'formatted_authors': formatted_authors,
                'pages': sorted(pages),
            })
        except Reference.DoesNotExist:
            bibliography.append({
                'number': ref_id,
                'reference': None,
                'ref_id': ref_id,
                'pages': [],
            })

    return bibliography


def get_filtered_user_answers(request, target_user):
    """Get filtered user answers based on POST options."""
    entry_ids = None
    order = 'oldest'
    root_question_id = None

    if request.method == 'POST':
        entry_ids_str = request.POST.get('entry_ids', '')
        if entry_ids_str:
            entry_ids = [int(id.strip()) for id in entry_ids_str.split(',') if id.strip()]
        order = request.POST.get('order', 'oldest')
        root_question_id = request.POST.get('root_question_id', '')

    user_answers = Answer.objects.filter(user=target_user)

    if root_question_id:
        try:
            root_id = int(root_question_id)
            question_ids = get_all_descendant_question_ids(root_id, target_user)
            user_answers = user_answers.filter(question_id__in=question_ids)
        except (ValueError, TypeError):
            pass

    if entry_ids:
        user_answers = user_answers.filter(id__in=entry_ids)

    if order == 'newest':
        user_answers = user_answers.order_by('-created_at')
    elif order == 'oldest':
        user_answers = user_answers.order_by('created_at')
    elif order == 'custom' and entry_ids:
        user_answers = list(user_answers)
        answers_dict = {ans.id: ans for ans in user_answers}
        user_answers = [answers_dict[id] for id in entry_ids if id in answers_dict]
    else:
        user_answers = user_answers.order_by('created_at')

    return user_answers


def is_custom_order_request(request):
    if request.method != 'POST':
        return False
    order = request.POST.get('order', 'oldest')
    entry_ids_str = (request.POST.get('entry_ids', '') or '').strip()
    return order == 'custom' and bool(entry_ids_str)


@login_required
def download_entries_json(request, username):
    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse({'error': 'Bu işlemi yapmaya yetkiniz yok.'}, status=403)

    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    questions_dict = {}
    for ans in user_answers:
        q = ans.question
        if q.id not in questions_dict:
            questions_dict[q.id] = {'question': q, 'answers': []}
        questions_dict[q.id]['answers'].append(ans)

    questions_data = []
    for _, q_data in questions_dict.items():
        q = q_data['question']
        q_answers = q_data['answers']

        answers_data = []
        for ans in q_answers:
            answers_data.append({
                'answer_text': ans.answer_text,
                'answer_created_at': ans.created_at.isoformat(),
                'answer_user': ans.user.username,
            })
        questions_data.append({
            'question_text': q.question_text,
            'question_created_at': q.created_at.isoformat(),
            'answers': answers_data,
        })

    entries_data = []
    if custom_order:
        for ans in user_answers:
            entries_data.append({
                'question_text': ans.question.question_text,
                'question_slug': ans.question.slug,
                'answer_id': ans.id,
                'answer_text': ans.answer_text,
                'answer_created_at': ans.created_at.isoformat(),
            })

    bibliography = collect_user_bibliography(target_user, user_answers)
    references_data = []
    for bib_item in bibliography:
        if bib_item.get('reference'):
            ref = bib_item['reference']
            ref_dict = {
                'number': bib_item['number'],
                'authors': bib_item['formatted_authors'],
                'year': ref.year,
                'title': ref.metin_ismi or '',
                'rest': ref.rest,
                'pages_used': bib_item['pages'],
            }
            if ref.abbreviation:
                ref_dict['abbreviation'] = ref.abbreviation
            references_data.append(ref_dict)
        else:
            references_data.append({
                'number': bib_item['number'],
                'ref_id': bib_item.get('ref_id'),
                'error': 'Reference not found',
            })

    final_data = {
        'username': target_user.username,
        'questions': questions_data,
        'entries': entries_data,
        'references': references_data,
    }

    json_string = json.dumps(final_data, ensure_ascii=False, indent=2)

    response = HttpResponse(json_string, content_type='application/json; charset=utf-8')
    response['Content-Disposition'] = 'attachment; filename="entries.json"'
    return response


@login_required
def download_entries_xlsx(request, username):
    from openpyxl import Workbook

    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse({'error': 'Bu işlemi yapmaya yetkiniz yok.'}, status=403)

    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    wb = Workbook()
    ws = wb.active
    ws.title = 'Entries'

    if custom_order:
        ws.cell(row=1, column=1, value='Soru')
        ws.cell(row=1, column=2, value='Tarih')
        ws.cell(row=1, column=3, value='Entry')
        for i, ans in enumerate(user_answers, start=2):
            ws.cell(row=i, column=1, value=ans.question.question_text)
            ws.cell(row=i, column=2, value=ans.created_at.strftime('%Y-%m-%d %H:%M'))
            ws.cell(row=i, column=3, value=ans.answer_text)
    else:
        questions_dict = {}
        for ans in user_answers:
            q = ans.question
            if q.id not in questions_dict:
                questions_dict[q.id] = {'question': q, 'answers': []}
            questions_dict[q.id]['answers'].append(ans)

        row_idx = 1
        for _, q_data in questions_dict.items():
            question = q_data['question']
            q_answers = q_data['answers']

            ws.cell(row=row_idx, column=1, value=question.question_text)

            answer_start_row = row_idx
            for j, ans in enumerate(q_answers):
                current_row = answer_start_row + j
                ws.cell(row=current_row, column=2, value=ans.answer_text)

            row_idx = answer_start_row + max(len(q_answers), 1)
            row_idx += 1

    bibliography = collect_user_bibliography(target_user, user_answers)
    if bibliography:
        ws_refs = wb.create_sheet(title='Kaynaklar')
        ws_refs.cell(row=1, column=1, value='No')
        ws_refs.cell(row=1, column=2, value='Yazarlar')
        ws_refs.cell(row=1, column=3, value='Yıl')
        ws_refs.cell(row=1, column=4, value='Metin İsmi')
        ws_refs.cell(row=1, column=5, value='Künye')
        ws_refs.cell(row=1, column=6, value='Kullanılan Sayfalar')

        for idx, bib_item in enumerate(bibliography, start=2):
            if bib_item.get('reference'):
                ref = bib_item['reference']
                ws_refs.cell(row=idx, column=1, value=bib_item['number'])
                ws_refs.cell(row=idx, column=2, value=bib_item['formatted_authors'])
                ws_refs.cell(row=idx, column=3, value=ref.year)
                ws_refs.cell(row=idx, column=4, value=ref.metin_ismi or '')
                ws_refs.cell(row=idx, column=5, value=ref.rest)
                ws_refs.cell(row=idx, column=6, value=', '.join(bib_item['pages']) if bib_item['pages'] else '')
            else:
                ws_refs.cell(row=idx, column=1, value=bib_item['number'])
                ws_refs.cell(row=idx, column=2, value='Kaynak bulunamadı')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="entries.xlsx"'

    wb.save(response)
    return response


def insert_toc(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_answer_text_to_docx(document, answer_text):
    if not answer_text:
        return

    normalized = str(answer_text).replace('\r\n', '\n').replace('\r', '\n').strip()
    if not normalized:
        return

    blocks = re.split(r'\n{2,}', normalized)
    for block in blocks:
        block = (block or '').strip('\n')
        if not block.strip():
            continue

        paragraph = document.add_paragraph()
        lines = block.split('\n')
        for idx, line in enumerate(lines):
            run = paragraph.add_run(line)
            if idx < len(lines) - 1:
                run.add_break()


def add_question_tree_to_docx(doc, question, target_user, level=1, visited=None):
    from docx.shared import Pt, RGBColor

    if visited is None:
        visited = set()
    if question.id in visited:
        return
    visited.add(question.id)

    doc.add_heading(question.question_text, level=level)

    user_answers = question.answers.filter(user=target_user).order_by('created_at')
    for answer in user_answers:
        date_str = answer.created_at.strftime('%Y-%m-%d %H:%M')
        p = doc.add_paragraph()
        run = p.add_run(date_str + '  ')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(140, 140, 140)
        run.italic = True
        add_answer_text_to_docx(doc, answer.answer_text)
        doc.add_paragraph('')

    subquestions_rels = QuestionRelationship.objects.filter(parent=question, user=target_user).select_related('child').order_by('created_at')
    for rel in subquestions_rels:
        add_question_tree_to_docx(doc, rel.child, target_user, level=min(level + 1, 9), visited=visited)


@login_required
def download_entries_docx(request, username):
    from docx import Document
    from docx.shared import Pt, RGBColor

    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse({'error': 'Bu işlemi yapmaya yetkiniz yok.'}, status=403)

    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    document = Document()
    document.add_heading(f'{target_user.username} Entries', 0)

    if custom_order:
        last_question_id = None
        for ans in user_answers:
            if ans.question_id != last_question_id:
                document.add_heading(ans.question.question_text, level=1)
                last_question_id = ans.question_id

            date_str = ans.created_at.strftime('%Y-%m-%d %H:%M')
            p = document.add_paragraph()
            run = p.add_run(date_str + '  ')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(140, 140, 140)
            run.italic = True
            add_answer_text_to_docx(document, ans.answer_text)
            document.add_paragraph('')
    else:
        questions_dict = {}
        for ans in user_answers:
            q = ans.question
            if q.id not in questions_dict:
                questions_dict[q.id] = {'question': q, 'answers': []}
            questions_dict[q.id]['answers'].append(ans)

        toc_paragraph = document.add_paragraph()
        insert_toc(toc_paragraph)

        instruction_text = (
            "Belgeyi açtıktan sonra içindekiler bölümünü görmek için, Word içerisinde "
            "alanı (veya tüm belgeyi) güncellemeniz gerekir (sağ tıklayıp 'Update Field' veya Ctrl+A ardından F9'a basabilirsiniz)."
        )
        document.add_paragraph(instruction_text)
        document.add_page_break()

        for _, q_data in questions_dict.items():
            question = q_data['question']
            q_answers = q_data['answers']

            document.add_heading(question.question_text, level=1)

            for answer in q_answers:
                date_str = answer.created_at.strftime('%Y-%m-%d %H:%M')
                p = document.add_paragraph()
                run = p.add_run(date_str + '  ')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(140, 140, 140)
                run.italic = True
                add_answer_text_to_docx(document, answer.answer_text)
                document.add_paragraph('')

    bibliography = collect_user_bibliography(target_user, user_answers)
    if bibliography:
        document.add_page_break()
        document.add_heading('Kaynakça', level=1)

        for bib_item in bibliography:
            if bib_item.get('reference'):
                ref = bib_item['reference']
                ref_text = f"[{bib_item['number']}] {bib_item['formatted_authors']} ({ref.year})"
                if ref.metin_ismi:
                    ref_text += f", {ref.metin_ismi}"
                ref_text += f", {ref.rest}"

                if bib_item['pages']:
                    ref_text += f" (Kullanılan sayfalar: {', '.join(bib_item['pages'])})"

                p = document.add_paragraph(ref_text)
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            else:
                ref_text = f"[{bib_item['number']}] Kaynak bulunamadı (ID: {bib_item.get('ref_id')})"
                p = document.add_paragraph(ref_text)
                p.paragraph_format.left_indent = Pt(18)

    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(
        f.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{target_user.username}_entries.docx"'
    return response


@login_required
def download_entries_pdf(request, username):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    target_user = get_object_or_404(User, username=username)
    if request.user != target_user and not request.user.is_superuser:
        return JsonResponse({'error': 'Bu işlemi yapmaya yetkiniz yok.'}, status=403)

    user_answers = get_filtered_user_answers(request, target_user)
    custom_order = is_custom_order_request(request)

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    elements = []

    try:
        import os
        from django.conf import settings
        import logging
        logger = logging.getLogger(__name__)

        font_dir = os.path.join(settings.BASE_DIR, 'static', 'fonts')
        dejavu_regular = os.path.join(font_dir, 'DejaVuSans.ttf')
        dejavu_bold = os.path.join(font_dir, 'DejaVuSans-Bold.ttf')

        logger.info(f'Font paths - Regular: {dejavu_regular}, Bold: {dejavu_bold}')
        logger.info(f'Regular exists: {os.path.exists(dejavu_regular)}, Bold exists: {os.path.exists(dejavu_bold)}')

        if os.path.exists(dejavu_regular) and os.path.exists(dejavu_bold):
            pdfmetrics.registerFont(TTFont('TurkishFont', dejavu_regular))
            pdfmetrics.registerFont(TTFont('TurkishFont-Bold', dejavu_bold))
            font_name = 'TurkishFont'
            font_name_bold = 'TurkishFont-Bold'
            logger.info('DejaVu Sans fonts successfully registered!')
        else:
            logger.warning('DejaVu fonts not found! Falling back to Helvetica (no Turkish support)')
            font_name = 'Helvetica'
            font_name_bold = 'Helvetica-Bold'
    except Exception as e:
        logger.error(f'Font registration failed: {e}')
        font_name = 'Helvetica'
        font_name_bold = 'Helvetica-Bold'

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'], fontName=font_name, fontSize=24,
        textColor='#2c3e50', spaceAfter=30, alignment=TA_CENTER
    )
    h1_style = ParagraphStyle(
        'CustomH1', parent=styles['Heading1'], fontName=font_name, fontSize=16,
        textColor='#2c3e50', spaceAfter=12, spaceBefore=12
    )
    answer_style = ParagraphStyle(
        'AnswerText', parent=styles['BodyText'], fontName=font_name, fontSize=10,
        textColor='#2c3e50', spaceAfter=12, leftIndent=20
    )
    date_style = ParagraphStyle(
        'DateStyle', parent=styles['Normal'], fontName=font_name, fontSize=8,
        textColor='#95a5a6', spaceAfter=6, leftIndent=20
    )

    title = Paragraph(f"{target_user.username} - Entry'ler", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.3 * inch))

    def clean_text(text):
        if not text:
            return ''
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('\n', '<br/>')
        return text

    if custom_order:
        last_question_id = None
        for answer in user_answers:
            if answer.question_id != last_question_id:
                elements.append(Paragraph(clean_text(answer.question.question_text), h1_style))
                elements.append(Spacer(1, 0.1 * inch))
                last_question_id = answer.question_id

            date_str = answer.created_at.strftime('%Y-%m-%d %H:%M')
            elements.append(Paragraph(f'<i>{date_str}</i>', date_style))
            elements.append(Paragraph(clean_text(answer.answer_text), answer_style))
            elements.append(Spacer(1, 0.15 * inch))
    else:
        questions_dict = {}
        for ans in user_answers:
            q = ans.question
            if q.id not in questions_dict:
                questions_dict[q.id] = {'question': q, 'answers': []}
            questions_dict[q.id]['answers'].append(ans)

        toc_style = ParagraphStyle(
            'TOCHeading', parent=styles['Heading1'], fontName=font_name_bold, fontSize=18,
            textColor='#2c3e50', spaceAfter=12
        )
        toc_item_style = ParagraphStyle(
            'TOCItem', parent=styles['Normal'], fontName=font_name, fontSize=11,
            textColor='#34495e', spaceAfter=6, leftIndent=20
        )

        elements.append(Paragraph('İçindekiler', toc_style))
        elements.append(Spacer(1, 0.2 * inch))

        for idx, (_, q_data) in enumerate(questions_dict.items(), 1):
            question = q_data['question']
            toc_text = f"{idx}. {clean_text(question.question_text[:100])}"
            if len(question.question_text) > 100:
                toc_text += '...'
            elements.append(Paragraph(toc_text, toc_item_style))

        elements.append(PageBreak())

        for _, q_data in questions_dict.items():
            question = q_data['question']
            q_answers = q_data['answers']

            elements.append(Paragraph(clean_text(question.question_text), h1_style))
            elements.append(Spacer(1, 0.1 * inch))

            for answer in q_answers:
                date_str = answer.created_at.strftime('%Y-%m-%d %H:%M')
                elements.append(Paragraph(f'<i>{date_str}</i>', date_style))
                elements.append(Paragraph(clean_text(answer.answer_text), answer_style))
                elements.append(Spacer(1, 0.15 * inch))

            elements.append(PageBreak())

    bibliography = collect_user_bibliography(target_user, user_answers)
    if bibliography:
        elements.append(PageBreak())

        bib_heading_style = ParagraphStyle(
            'BibliographyHeading', parent=styles['Heading1'], fontName=font_name_bold,
            fontSize=18, textColor='#2c3e50', spaceAfter=20, spaceBefore=12
        )
        bib_item_style = ParagraphStyle(
            'BibliographyItem', parent=styles['BodyText'], fontName=font_name, fontSize=10,
            textColor='#2c3e50', spaceAfter=10, leftIndent=20, firstLineIndent=-20
        )

        elements.append(Paragraph('Kaynakça', bib_heading_style))
        elements.append(Spacer(1, 0.2 * inch))

        for bib_item in bibliography:
            if bib_item.get('reference'):
                ref = bib_item['reference']
                ref_text = f"[{bib_item['number']}] {clean_text(bib_item['formatted_authors'])} ({ref.year})"
                if ref.metin_ismi:
                    ref_text += f", {clean_text(ref.metin_ismi)}"
                ref_text += f", {clean_text(ref.rest)}"
                if bib_item['pages']:
                    pages_str = ', '.join(bib_item['pages'])
                    ref_text += f" (Kullanılan sayfalar: {pages_str})"
                elements.append(Paragraph(ref_text, bib_item_style))
            else:
                ref_text = f"[{bib_item['number']}] Kaynak bulunamadı (ID: {bib_item.get('ref_id')})"
                elements.append(Paragraph(ref_text, bib_item_style))

    doc.build(elements)
    buffer.seek(0)
    pdf_content = buffer.read()

    response = HttpResponse(pdf_content, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{target_user.username}_entries.pdf"'
    return response


@login_required
def filter_answers(request, slug):
    question = get_object_or_404(Question, slug=slug)

    my_answers = request.GET.get('my_answers')
    followed = request.GET.get('followed')
    username = request.GET.get('username', '').strip()
    keyword = request.GET.get('keyword', '').strip()

    answers = question.answers.all().order_by('created_at')

    if my_answers == 'on':
        answers = answers.filter(user=request.user)

    if followed == 'on':
        user_profile = request.user.userprofile
        followed_profiles = user_profile.following.all()
        followed_users = [p.user for p in followed_profiles]
        answers = answers.filter(user__in=followed_users)

    if username:
        answers = answers.filter(user__username__icontains=username)

    if keyword:
        answers = answers.filter(answer_text__icontains=keyword)

    content_type_answer = ContentType.objects.get_for_model(Answer)
    answer_ids = answers.values_list('id', flat=True)

    saved_items = SavedItem.objects.filter(content_type=content_type_answer, object_id__in=answer_ids).values('object_id').annotate(count=Count('id'))
    answer_save_dict = {item['object_id']: item['count'] for item in saved_items}

    saved_answer_ids = SavedItem.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values_list('object_id', flat=True)

    user_votes = Vote.objects.filter(
        user=request.user,
        content_type=content_type_answer,
        object_id__in=answer_ids
    ).values('object_id', 'value')
    user_vote_dict = {v['object_id']: v['value'] for v in user_votes}

    answers = list(answers.select_related('user', 'question'))
    attach_answer_revision_metadata(answers, current_user=request.user)
    for ans in answers:
        ans.user_vote_value = user_vote_dict.get(ans.id, 0)
        ans.upvotes = Vote.objects.filter(object_id=ans.id, value=1).count()
        ans.downvotes = Vote.objects.filter(object_id=ans.id, value=-1).count()

    html_content = render_to_string(
        'core/_answers_list.html',
        {
            'answers': answers,
            'question': question,
            'answer_save_dict': answer_save_dict,
            'saved_answer_ids': saved_answer_ids,
            'search_keyword': keyword,
        },
        request=request,
    )
    return HttpResponse(html_content)
