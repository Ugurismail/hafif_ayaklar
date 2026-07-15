"""Academic DOCX export for selected user entries."""

from __future__ import annotations

import re
import zipfile
from collections import deque
from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree

from .models import QuestionRelationship, Reference


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKGREL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

NS = {"w": W_NS, "r": R_NS, "rel": PKGREL_NS, "ct": CT_NS}
REL_TYPE_FOOTNOTES = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)
CT_FOOTNOTES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
)

CITATION_RE = re.compile(
    r"\((?:kaynak|k)\s*:\s*(?P<reference_id>\d+)"
    r"(?:\s*,?\s*(?:sayfa|s)\s*:\s*(?P<page>[^)]+?))?\s*\)",
    re.IGNORECASE,
)
SPOILER_RE = re.compile(
    r"-g-\s*(?P<new>.*?)\s*-g-|--gizli--(?P<old>.*?)--gizli--",
    re.IGNORECASE | re.DOTALL,
)
INLINE_RE = re.compile(
    r"(?P<footnote>\[\[PAPER_FOOTNOTE_\d+\]\])"
    r"|(?P<link>\[(?P<link_text>[^\]]+)\]\((?P<link_url>https?://[^)]+)\))"
    r"|(?P<bold>\*\*(?P<bold_text>.+?)\*\*)"
    r"|(?P<italic>(?<!\*)\*(?P<italic_text>[^*\n]+?)\*(?!\*))"
)
TURKISH_ALPHABET = "abcçdefgğhıijklmnoöprsştuüvyz"
TURKISH_SORT_ORDER = {
    character: index
    for index, character in enumerate(TURKISH_ALPHABET)
}


def _word_tag(local_name):
    return f"{{{W_NS}}}{local_name}"


def _xml_bytes(root):
    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def _turkish_sort_key(value):
    normalized = (value or "").translate(
        str.maketrans({"I": "ı", "İ": "i"})
    ).lower()
    unknown_offset = len(TURKISH_SORT_ORDER)
    return tuple(
        TURKISH_SORT_ORDER.get(character, unknown_offset + ord(character))
        for character in normalized
    )


def _author_parts(reference):
    surnames = [
        value.strip()
        for value in (reference.author_surname or "").split(";")
        if value.strip()
    ]
    names = [
        value.strip()
        for value in (reference.author_name or "").split(";")
        if value.strip()
    ]
    author_count = max(len(surnames), len(names))
    authors = []
    for index in range(author_count):
        surname = surnames[index] if index < len(surnames) else ""
        name = names[index] if index < len(names) else ""
        if surname or name:
            authors.append((surname, name))
    return authors


def _citation_author(reference):
    authors = _author_parts(reference)
    if not authors:
        return "Anonim"
    first_surname, first_name = authors[0]
    display = first_surname or first_name or "Anonim"
    if len(authors) > 1:
        display += " vd."
    return display


def _bibliography_author(reference):
    authors = _author_parts(reference)
    if not authors:
        return "Anonim"
    formatted = []
    for surname, name in authors:
        if surname and name:
            formatted.append(f"{surname}, {name}")
        else:
            formatted.append(surname or name)
    return "; ".join(formatted)


def _bibliography_sort_value(reference):
    authors = _author_parts(reference)
    if not authors:
        return "anonim"
    surname, name = authors[0]
    return _turkish_sort_key(surname or name)


def _next_relationship_id(rels_root):
    highest = 0
    for relationship in rels_root.findall(f"{{{PKGREL_NS}}}Relationship"):
        match = re.match(r"rId(\d+)$", relationship.get("Id") or "")
        if match:
            highest = max(highest, int(match.group(1)))
    return f"rId{highest + 1}"


def _ensure_footnote_relationship(rels_root):
    for relationship in rels_root.findall(f"{{{PKGREL_NS}}}Relationship"):
        if relationship.get("Type") == REL_TYPE_FOOTNOTES:
            return False
    relationship = etree.SubElement(
        rels_root,
        f"{{{PKGREL_NS}}}Relationship",
    )
    relationship.set("Id", _next_relationship_id(rels_root))
    relationship.set("Type", REL_TYPE_FOOTNOTES)
    relationship.set("Target", "footnotes.xml")
    return True


def _ensure_footnote_content_type(content_types_root):
    part_name = "/word/footnotes.xml"
    for override in content_types_root.findall(f"{{{CT_NS}}}Override"):
        if override.get("PartName") == part_name:
            override.set("ContentType", CT_FOOTNOTES)
            return False
    override = etree.SubElement(content_types_root, f"{{{CT_NS}}}Override")
    override.set("PartName", part_name)
    override.set("ContentType", CT_FOOTNOTES)
    return True


def _append_direct_run(paragraph, text, *, superscript=False, bold=False):
    run = etree.SubElement(paragraph, _word_tag("r"))
    run_properties = etree.SubElement(run, _word_tag("rPr"))
    fonts = etree.SubElement(run_properties, _word_tag("rFonts"))
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(_word_tag(attribute), "Times New Roman")
    size = etree.SubElement(run_properties, _word_tag("sz"))
    size.set(_word_tag("val"), "18")
    size_cs = etree.SubElement(run_properties, _word_tag("szCs"))
    size_cs.set(_word_tag("val"), "18")
    if superscript:
        vertical_align = etree.SubElement(run_properties, _word_tag("vertAlign"))
        vertical_align.set(_word_tag("val"), "superscript")
    if bold:
        etree.SubElement(run_properties, _word_tag("b"))
    text_element = etree.SubElement(run, _word_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    return run


def _make_footnotes_part(notes):
    root = etree.Element(_word_tag("footnotes"), nsmap={"w": W_NS, "r": R_NS})

    for note_id, separator_type in ((-1, "separator"), (0, "continuationSeparator")):
        note = etree.SubElement(root, _word_tag("footnote"))
        note.set(_word_tag("id"), str(note_id))
        note.set(_word_tag("type"), separator_type)
        paragraph = etree.SubElement(note, _word_tag("p"))
        run = etree.SubElement(paragraph, _word_tag("r"))
        etree.SubElement(run, _word_tag(separator_type))

    for note in notes:
        footnote = etree.SubElement(root, _word_tag("footnote"))
        footnote.set(_word_tag("id"), str(note["id"]))

        paragraph = etree.SubElement(footnote, _word_tag("p"))
        paragraph_properties = etree.SubElement(paragraph, _word_tag("pPr"))
        spacing = etree.SubElement(paragraph_properties, _word_tag("spacing"))
        spacing.set(_word_tag("after"), "0")
        spacing.set(_word_tag("line"), "240")
        spacing.set(_word_tag("lineRule"), "auto")

        reference_run = etree.SubElement(paragraph, _word_tag("r"))
        reference_properties = etree.SubElement(reference_run, _word_tag("rPr"))
        reference_style = etree.SubElement(reference_properties, _word_tag("rStyle"))
        reference_style.set(_word_tag("val"), "FootnoteReference")
        etree.SubElement(reference_run, _word_tag("footnoteRef"))
        _append_direct_run(
            paragraph,
            note["mark"],
            superscript=True,
            bold=True,
        )
        _append_direct_run(paragraph, f" {note['text']}")

    return root


def _replace_footnote_marker(document_root, marker, note_id, mark):
    for text_element in document_root.xpath(".//w:t", namespaces=NS):
        if marker not in (text_element.text or ""):
            continue

        text_element.text = text_element.text.replace(marker, "")
        source_run = text_element.getparent()
        while source_run is not None and source_run.tag != _word_tag("r"):
            source_run = source_run.getparent()
        if source_run is None:
            return False

        parent = source_run.getparent()
        reference_run = etree.Element(_word_tag("r"))
        reference_properties = etree.SubElement(reference_run, _word_tag("rPr"))
        reference_style = etree.SubElement(reference_properties, _word_tag("rStyle"))
        reference_style.set(_word_tag("val"), "FootnoteReference")
        reference = etree.SubElement(reference_run, _word_tag("footnoteReference"))
        reference.set(_word_tag("id"), str(note_id))
        reference.set(_word_tag("customMarkFollows"), "1")

        custom_marker_run = etree.Element(_word_tag("r"))
        custom_marker_properties = etree.SubElement(custom_marker_run, _word_tag("rPr"))
        custom_marker_style = etree.SubElement(custom_marker_properties, _word_tag("rStyle"))
        custom_marker_style.set(_word_tag("val"), "FootnoteReference")
        custom_marker = etree.SubElement(custom_marker_run, _word_tag("t"))
        custom_marker.text = "\u200b"

        visible_mark_run = etree.Element(_word_tag("r"))
        visible_mark_properties = etree.SubElement(visible_mark_run, _word_tag("rPr"))
        visible_mark_style = etree.SubElement(visible_mark_properties, _word_tag("rStyle"))
        visible_mark_style.set(_word_tag("val"), "FootnoteReference")
        visible_mark = etree.SubElement(visible_mark_run, _word_tag("t"))
        visible_mark.text = mark
        source_index = parent.index(source_run)
        parent.insert(source_index + 1, reference_run)
        parent.insert(source_index + 2, custom_marker_run)
        parent.insert(source_index + 3, visible_mark_run)
        return True
    return False


def _patch_footnotes(docx_bytes, notes):
    if not notes:
        return docx_bytes

    input_buffer = BytesIO(docx_bytes)
    output_buffer = BytesIO()
    with zipfile.ZipFile(input_buffer, "r") as source:
        document_root = etree.fromstring(source.read("word/document.xml"))
        for note in notes:
            if not _replace_footnote_marker(
                document_root,
                note["marker"],
                note["id"],
                note["mark"],
            ):
                raise ValueError(f"Dipnot işareti belgede bulunamadı: {note['marker']}")

        relationships_root = etree.fromstring(
            source.read("word/_rels/document.xml.rels")
        )
        content_types_root = etree.fromstring(source.read("[Content_Types].xml"))
        _ensure_footnote_relationship(relationships_root)
        _ensure_footnote_content_type(content_types_root)

        replacements = {
            "word/document.xml": _xml_bytes(document_root),
            "word/_rels/document.xml.rels": _xml_bytes(relationships_root),
            "[Content_Types].xml": _xml_bytes(content_types_root),
            "word/footnotes.xml": _xml_bytes(_make_footnotes_part(notes)),
        }
        existing_names = set(source.namelist())

        with zipfile.ZipFile(output_buffer, "w", zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                if info.filename in replacements:
                    target.writestr(info, replacements[info.filename])
                else:
                    target.writestr(info, source.read(info.filename))
            for name, content in replacements.items():
                if name not in existing_names:
                    target.writestr(name, content)

    return output_buffer.getvalue()


def _set_run_font(run, *, size=None, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Times New Roman")
    run_properties.append(fonts)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_internal_hyperlink(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Times New Roman")
    run_properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_bookmark(paragraph, bookmark_id, bookmark_name):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_inline_text(paragraph, text):
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            _set_run_font(run)

        if match.group("footnote"):
            run = paragraph.add_run(match.group("footnote"))
            _set_run_font(run)
        elif match.group("link"):
            _add_hyperlink(paragraph, match.group("link_text"), match.group("link_url"))
        elif match.group("bold"):
            run = paragraph.add_run(match.group("bold_text"))
            _set_run_font(run, bold=True)
        elif match.group("italic"):
            run = paragraph.add_run(match.group("italic_text"))
            _set_run_font(run, italic=True)
        position = match.end()

    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_run_font(run)


def _add_horizontal_rule(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7B7B7")
    borders.append(bottom)
    properties.append(borders)


def _add_answer_body(document, text, renderer, question_level):
    prepared = renderer.prepare_text(text)
    if not prepared.strip():
        return

    blocks = re.split(r"\n{2,}", prepared.strip())
    for block in blocks:
        block = block.strip("\n")
        if not block.strip():
            continue
        if block.strip() in {"--", "---", "***"}:
            _add_horizontal_rule(document)
            continue

        heading_match = re.fullmatch(r"\s*(#{1,6})\s+(.+?)\s*#*\s*", block)
        if heading_match:
            relative_level = max(1, (len(heading_match.group(1)) + 1) // 2)
            level = min(question_level + relative_level, 9)
            heading = document.add_heading(level=level)
            _add_inline_text(heading, heading_match.group(2))
            continue

        paragraph = document.add_paragraph()
        if all(line.lstrip().startswith(">") for line in block.splitlines()):
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.right_indent = Cm(0.5)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            lines = [line.lstrip()[1:].lstrip() for line in block.splitlines()]
            block = "\n".join(lines)
            paragraph.style = document.styles["Paper Quote"]

        for index, line in enumerate(block.split("\n")):
            _add_inline_text(paragraph, line)
            if index < len(block.split("\n")) - 1:
                paragraph.add_run().add_break(WD_BREAK.LINE)


def _insert_toc(document, outline):
    heading = document.add_paragraph(style="Paper TOC Heading")
    run = heading.add_run("İçindekiler")
    _set_run_font(run, size=18, bold=True)

    for item in outline:
        paragraph = document.add_paragraph(style="Paper TOC Entry")
        paragraph.paragraph_format.left_indent = Cm((item["level"] - 1) * 0.65)
        _add_internal_hyperlink(
            paragraph,
            item["title"],
            item["bookmark_name"],
        )
    document.add_page_break()


def _add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    _set_run_font(run, size=10)


def _configure_document(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    heading_sizes = {
        1: 16,
        2: 14,
        3: 12,
        4: 11,
        5: 11,
        6: 10,
        7: 10,
        8: 10,
        9: 10,
    }
    for level, size in heading_sizes.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(16 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(8 if level == 1 else 5)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    toc_heading = document.styles.add_style("Paper TOC Heading", WD_STYLE_TYPE.PARAGRAPH)
    toc_heading.font.name = "Times New Roman"
    toc_heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    toc_heading.font.size = Pt(18)
    toc_heading.font.bold = True
    toc_heading.font.color.rgb = RGBColor(0, 0, 0)
    toc_heading.paragraph_format.space_after = Pt(16)

    toc_entry = document.styles.add_style("Paper TOC Entry", WD_STYLE_TYPE.PARAGRAPH)
    toc_entry.font.name = "Times New Roman"
    toc_entry._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    toc_entry.font.size = Pt(12)
    toc_entry.font.color.rgb = RGBColor(0, 0, 0)
    toc_entry.paragraph_format.space_after = Pt(6)
    toc_entry.paragraph_format.keep_together = True

    quote = document.styles.add_style("Paper Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote.base_style = normal
    quote.font.name = "Times New Roman"
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.paragraph_format.line_spacing = 1.15
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(8)

    bibliography = document.styles.add_style("Paper Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    bibliography.base_style = normal
    bibliography.font.name = "Times New Roman"
    bibliography.font.size = Pt(11)
    bibliography.paragraph_format.left_indent = Cm(1.25)
    bibliography.paragraph_format.first_line_indent = Cm(-1.25)
    bibliography.paragraph_format.line_spacing = 1.15
    bibliography.paragraph_format.space_after = Pt(8)

    footnote_reference = document.styles.add_style(
        "Footnote Reference",
        WD_STYLE_TYPE.CHARACTER,
    )
    footnote_reference.font.name = "Times New Roman"
    footnote_reference._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Times New Roman",
    )
    footnote_reference.font.size = Pt(9)
    footnote_reference.font.superscript = True

    _add_page_number(section)


def _question_levels(answers, target_user, root_question_id=None):
    question_ids = {answer.question_id for answer in answers}
    if not question_ids:
        return {}

    relationships = list(
        QuestionRelationship.objects.filter(user=target_user).values_list(
            "parent_id",
            "child_id",
        )
    )
    children = {}
    parents = {}
    graph_nodes = set(question_ids)
    for parent_id, child_id in relationships:
        children.setdefault(parent_id, []).append(child_id)
        parents.setdefault(child_id, []).append(parent_id)
        graph_nodes.update((parent_id, child_id))

    levels = {}
    if root_question_id:
        queue = deque([(root_question_id, 1)])
        visited = set()
        while queue:
            question_id, level = queue.popleft()
            if question_id in visited:
                continue
            visited.add(question_id)
            levels[question_id] = min(level, 9)
            for child_id in children.get(question_id, []):
                queue.append((child_id, level + 1))

    roots = [node_id for node_id in graph_nodes if not parents.get(node_id)]
    queue = deque((root_id, 1) for root_id in roots)
    while queue:
        question_id, level = queue.popleft()
        existing = levels.get(question_id)
        if existing is not None and existing <= level:
            continue
        levels[question_id] = min(level, 9)
        for child_id in children.get(question_id, []):
            queue.append((child_id, level + 1))

    return {question_id: levels.get(question_id, 1) for question_id in question_ids}


class PaperTextRenderer:
    def __init__(self, answers):
        all_text = "\n".join(answer.answer_text or "" for answer in answers)
        reference_ids = {
            int(match.group("reference_id"))
            for match in CITATION_RE.finditer(all_text)
        }
        self.references = Reference.objects.in_bulk(reference_ids)
        self.used_reference_ids = set(reference_ids)
        self.notes = []

    def _replace_citation(self, match):
        reference_id = int(match.group("reference_id"))
        self.used_reference_ids.add(reference_id)
        reference = self.references.get(reference_id)
        if reference is None:
            return f"(Kaynak {reference_id})"

        parts = [_citation_author(reference), str(reference.year)]
        page = (match.group("page") or "").strip()
        if page:
            parts.append(f"s. {page}")
        return f"({', '.join(parts)})"

    def _clean_footnote_text(self, text):
        text = CITATION_RE.sub(self._replace_citation, text or "")
        text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", text, flags=re.DOTALL)
        return re.sub(r"\s+", " ", text).strip()

    def prepare_text(self, text):
        def replace_spoiler(match):
            note_text = (
                match.group("new")
                if match.group("new") is not None
                else match.group("old")
            )
            note_id = len(self.notes) + 1
            marker = f"[[PAPER_FOOTNOTE_{note_id}]]"
            self.notes.append(
                {
                    "id": note_id,
                    "marker": marker,
                    "mark": "*" * (((note_id - 1) % 3) + 1),
                    "text": self._clean_footnote_text(note_text),
                }
            )
            return marker

        prepared = SPOILER_RE.sub(replace_spoiler, str(text or ""))
        prepared = re.sub(
            r"[ \t]+(?=\[\[PAPER_FOOTNOTE_\d+\]\])",
            "\u00a0",
            prepared,
        )
        return CITATION_RE.sub(self._replace_citation, prepared)

    def bibliography(self):
        found = [
            self.references[reference_id]
            for reference_id in self.used_reference_ids
            if reference_id in self.references
        ]
        return sorted(found, key=_bibliography_sort_value)

    def missing_reference_ids(self):
        return sorted(self.used_reference_ids - set(self.references))


def _add_bibliography_entry(document, reference):
    paragraph = document.add_paragraph(style="Paper Bibliography")
    author_run = paragraph.add_run(
        f"{_bibliography_author(reference)} ({reference.year}). "
    )
    _set_run_font(author_run, size=11)
    if reference.metin_ismi:
        title_run = paragraph.add_run(reference.metin_ismi.rstrip(". "))
        _set_run_font(title_run, size=11, italic=True)
        punctuation_run = paragraph.add_run(". ")
        _set_run_font(punctuation_run, size=11)
    if reference.rest:
        rest_run = paragraph.add_run(reference.rest.strip())
        _set_run_font(rest_run, size=11)


def build_paper_docx(answers, target_user, root_question_id=None):
    """Build a styled academic DOCX and return its bytes."""
    answers = list(answers)
    document = Document()
    _configure_document(document)
    renderer = PaperTextRenderer(answers)
    levels = _question_levels(answers, target_user, root_question_id)
    bibliography = renderer.bibliography()
    missing_reference_ids = renderer.missing_reference_ids()

    outline = []
    bookmark_by_question_id = {}
    for answer in answers:
        if answer.question_id in bookmark_by_question_id:
            continue
        bookmark_name = f"paper_heading_{len(bookmark_by_question_id) + 1}"
        bookmark_by_question_id[answer.question_id] = bookmark_name
        outline.append(
            {
                "title": answer.question.question_text,
                "level": levels.get(answer.question_id, 1),
                "bookmark_name": bookmark_name,
            }
        )
    if bibliography or missing_reference_ids:
        outline.append(
            {
                "title": "Kaynakça",
                "level": 1,
                "bookmark_name": "paper_bibliography",
            }
        )

    document.core_properties.title = "Paper"
    document.core_properties.author = target_user.username
    _insert_toc(document, outline)

    last_question_id = None
    bookmarked_question_ids = set()
    for answer in answers:
        level = levels.get(answer.question_id, 1)
        if answer.question_id != last_question_id:
            heading = document.add_heading(answer.question.question_text, level=level)
            if answer.question_id not in bookmarked_question_ids:
                _add_bookmark(
                    heading,
                    len(bookmarked_question_ids) + 1,
                    bookmark_by_question_id[answer.question_id],
                )
                bookmarked_question_ids.add(answer.question_id)
            last_question_id = answer.question_id
        _add_answer_body(document, answer.answer_text, renderer, level)

    if bibliography or missing_reference_ids:
        document.add_page_break()
        bibliography_heading = document.add_heading("Kaynakça", level=1)
        _add_bookmark(
            bibliography_heading,
            len(bookmarked_question_ids) + 1,
            "paper_bibliography",
        )
        for reference in bibliography:
            _add_bibliography_entry(document, reference)
        for reference_id in missing_reference_ids:
            paragraph = document.add_paragraph(style="Paper Bibliography")
            run = paragraph.add_run(f"Kaynak kaydı bulunamadı (ID: {reference_id}).")
            _set_run_font(run, size=11)

    buffer = BytesIO()
    document.save(buffer)
    return _patch_footnotes(buffer.getvalue(), renderer.notes)
